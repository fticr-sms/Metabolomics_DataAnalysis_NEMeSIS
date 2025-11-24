
## Needed imports
import pandas as pd
import numpy as np
import panel as pn
import param
import holoviews as hv
import pyopenms as oms
from io import BytesIO
import plotly.express as px
import matplotlib.pyplot as plt
from metabolinks import read_data_from_xcel, align
from metabolinks.peak_alignment import alignment_summary

import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os

# Activating extensions
pn.extension('plotly', 'floatpanel', 'katex', 'ipywidgets', notifications=True)
hv.extension('plotly')
pn.config.sizing_mode="stretch_width"

download_icon = '''<svg xmlns="http://www.w3.org/2000/svg" class="icon icon-tabler icon-tabler-download" width="24" height="24" viewBox="0 0 24 24" stroke-width="2" stroke="currentColor" fill="none" stroke-linecap="round" stroke-linejoin="round">
   <path stroke="none" d="M0 0h24v24H0z" fill="none"></path>
   <path d="M4 17v2a2 2 0 0 0 2 2h12a2 2 0 0 0 2 -2v-2"></path>
   <path d="M7 11l5 5l5 -5"></path>
   <path d="M12 4l0 12"></path>
</svg>'''

# Important Text
alignment_description_string = """
<p>This Data Alignment software is geared towards <strong>Direct Infusion mass spectrometry data</strong>. That is,
metabolic features will be aligned based on their mass values alone, and it does not accept other parameters such as
retention time or collision cross-section.</p>
<p>The input can be either samples represented as lists of <em>m/z</em> peaks (as an Excel with a sample per sheet with the mass
peak list in each or as a csv file per sample) or as spectral raw data in mzML format.</p>
<p>To perform data alignment from <em>m/z</em> lists with a single <strong>Excel</strong> file (.xlsx or .xls), this
file should follow the format: <strong>one sample per Excel sheet</strong>; the <strong>name</strong> of the Excel
sheet should correspond to the <strong>sample name</strong>; each sheet should have in its <strong>first column the mass
values</strong> (<em>m/z</em>, neutral mass or equivalent) and in its <strong>second column&nbsp;</strong>the corresponding
<strong>intensity values</strong>; and finally, the <strong>first row</strong> should have the <strong>name of the two
columns</strong>, for example, &apos;<em>m/z</em>&apos; and &apos;I&apos; (this name should be consistent between samples).
The example file &apos;example_samples_to_align.xlsx&apos; is available in the Files_To_Align folder as guidance.</p>
<p>To perform data alignment from <em>m/z</em> lists from a set of <strong>csv</strong> files, each file should have the
following format: <strong>one sample per csv</strong> where the first column has the mass values</strong> (<em>m/z</em>,
neutral mass or equivalent) and the <strong>second column&nbsp;</strong>the corresponding <strong>intensity values</strong>.
The <strong>first row</strong> should have the name of the <em>m/z</em> column (usually m/z, this name should be consistent
between samples) and <strong>the second the name of the sample</strong>. The example file &apos;placeholder.csv&apos; is
available in the Files_To_Align folder as guidance.</p>
<p>To use spectral raw data in mzML format, <strong>multiple mzML files</strong> are expected. Converting spectral raw data into
the open mzML format can be done using freely available tools such as Proteowizard's MSConvert
(<a href=https://proteowizard.sourceforge.io/download.html target="_blank" rel="nofollow">https://proteowizard.sourceforge.io/download.html</a>).
</p>"""


mzML_conversion_specific_string = """
<p>Here, raw mzML spectral data (open mzML format) can be converted into <em>m/z</em> peak lists so data alignment can be
performed. For this purpose, the pyopenms Python package is used and extensive documentation is present in
<a href=https://openms.de/current_doxygen/html/classOpenMS_1_1PeakPickerHiRes.html target="_blank"
rel="nofollow">https://openms.de/current_doxygen/html/classOpenMS_1_1PeakPickerHiRes.html</a>. The <strong>signal to noise
threshold</strong> is the most critical parameter, the others shown may also be customized - see effects in the
aforementioned link. The description from the link is also near every parameter with a tooltip. Credit to pyopenms devs.
</p>
"""


data_alignment_specific_string = """
<p>The data alignment performed is made using the <strong><em>align</em></strong> function from the Metabolinks Python package.
It uses 2 parameters: the <strong>PPM Deviation Tolerance</strong> that defines the maximum tolerance (in parts per million)
to group metabolic features from different samples together and the <strong>Minimum Sample Number Appearance</strong> that
defines the minimum number of samples a metabolic feature must appear in to be kept in the final aligned Dataset.</p>
<p>After alignment, a short description of the alignment process will be shown as well as the aligned dataset (if this
aligned dataset is large, only the first 10,000 metabolic features will be shown). The dataset and an abridged description
of the metabolic feature alignments can be saved. The aligned dataset can be directly used in the NEMeSIS software.</p>"""


spectra_visualization_specific_string = """
<p>This section <strong>only</strong> appers if the aligned data comes from spectral mzML raw data. The objective is to
visualize the quality of the spectral data processing by observing the <strong>raw, centroided (after spectral processing) and
aligned (after data alignment) spectra</strong> in interactive graphs of a <strong>few samples (3-4)</strong> simultaneously.</p>"""



class AlignmentPage:
    def __init__(self):
        self.content = pn.Column("# Data Alignment Section", alignment_description_string, alignment_page)
    def view(self):
        return self.content


# Contains relevant parameters regarding Data Alignment
class ReadingMZML_Files(param.Parameterized):
    "Class to contain all relevant parameters regarding mzML File conversion."

    # Samples in the list
    samples = param.Dict()
    files = param.List()

    # Parameters for Raw Spectral Data Conversion
    signal_to_noise = param.Number(default=3.3) # 0 float
    spacing_difference_gap = param.Number(default=4.0) # Min 0 float
    spacing_difference = param.Number(default=1.5) # 0 float
    missing = param.Number(default=1) # 0 int
    signal_to_noise_win_len = param.Number(default=200.0) # 1 float
    signal_to_noise_bin_count = param.Number(default=30) # 3 int
    signal_to_noise_min_required_elements = param.Number(default=10) # 1 int
    #signal_to_noise_auto_max_stdev_factor = param.Number(default=3) # 0 to 999 float
    #signal_to_noise_auto_max_percentile = param.Number(default=95) # 0 to 100 int
    #signal_to_noise_auto_mode = param.Number(default=0) # -1, 0 or 1 int
    confirm_spectral_conv_button = param.Boolean(default=False)
    save_conversion_res_button = param.Boolean(default=False)

    # Spectral Raw Data Conversion Description
    desc = param.DataFrame()
    aligned_desc = param.String(default='')
    save_parameters = param.Dict()


    def _confirm_spectral_conv_button(self, event):
        "Performs conversion of spetral mzML data to peak lists and updates layout."

        # Updating layout
        while len(data_mzML_section)>3:
            data_mzML_section.pop(-1)

        # Loading Widget while spectra raw data conversion is happening
        data_mzML_section.append(pn.indicators.LoadingSpinner(value=True, size=120,
                                                        name='Performing Spectral mzML Raw Data Conversion...'))

        self.samples = {}

        # Go through every filename selected
        for f in filename.value:
            # Files
            f_name = f.split('\\')[-1].split('.')[0]

            # Start pyopenms
            exp = oms.MSExperiment()
            oms.MzMLFile().load(f, exp)
            spec = exp[0]
            mz, intensity = spec.get_peaks()
            raw_df = pd.DataFrame(data={f_name: intensity}, index= mz)

            centroided_spectra = oms.MSExperiment()
            cnt = oms.PeakPickerHiRes()

            # Set all parameters
            param = cnt.getParameters()
            param.setValue("signal_to_noise", self.signal_to_noise)
            param.setValue("spacing_difference_gap", self.spacing_difference_gap)
            param.setValue("spacing_difference", self.spacing_difference)
            param.setValue("missing", self.missing)
            param.setValue("SignalToNoise:win_len", self.signal_to_noise_win_len)
            param.setValue("SignalToNoise:bin_count", self.signal_to_noise_bin_count)
            param.setValue("SignalToNoise:min_required_elements", self.signal_to_noise_min_required_elements)
            cnt.setParameters(param)
            # Perform the conversion from mzML file to list of peaks
            cnt.pickExperiment(exp, centroided_spectra, True,)
            cent_mz, cent_ints = centroided_spectra[0].get_peaks()
            cent_df = pd.DataFrame(data={f_name: cent_ints}, index = cent_mz)

            # Description of Conversion
            data_mzML_section.insert(-1, f'{f_name} | Raw signals: {len(raw_df)} | Centroided signals: {len(cent_df)}')
            self.aligned_desc += f'{f_name} | Raw signals: {len(raw_df)} | Centroided signals: {len(cent_df)} \n'

            self.samples[f_name] = {'raw': raw_df, 'centroided': cent_df}

        # Savee Parameters Used
        self.save_parameters = {"signal_to_noise": self.signal_to_noise,
                                "spacing_difference_gap": self.spacing_difference_gap,
                                "spacing_difference": self.spacing_difference,
                                "missing": self.missing,
                                "signal_to_noise_win_len": self.signal_to_noise_win_len,
                                "signal_to_noise_bin_count": self.signal_to_noise_bin_count,
                                "signal_to_noise_min_required_elements": self.signal_to_noise_min_required_elements,}

        # Prepare Spectra Visualization
        spectra_store.update_widgets()

        # Remove Loading Spinner
        data_mzML_section.pop(-1)
        data_mzML_section.append(self.controls2)
        pn.state.notifications.success(f'Completed mzML file conversion.')


    def _save_conversion_res_button(self, event):
        "Save conversion results and perform data alignment."

        # Creating the Word Document
        document = docx.Document()

        # Adding the title
        document.add_heading('Spectral Raw Data Conversion to Peak Lists', 0)

        # Files used
        document.add_paragraph('The following spectral mzML raw data files were converted to mass peak lists:')
        for f in self.files:
            small_file = f.split('\\')[-1]
            document.add_paragraph(f'{small_file}', style='List Bullet')
        document.add_paragraph()

        # Parameters used
        document.add_paragraph('The following parameters were used with pyopenms for the conversion:')
        for k,v in self.save_parameters.items():
            document.add_paragraph(f'{k}: {v}', style='List Bullet')
        document.add_paragraph('See parameters details in https://openms.de/current_doxygen/html/classOpenMS_1_1PeakPickerHiRes.html')
        document.add_paragraph()

        # Results
        document.add_paragraph('The results of the conversion can be described by the following text:')
        res = self.aligned_desc.split('\n')
        for l in res:
            if l != '':
                document.add_paragraph(l, style='List Bullet')
        document.add_paragraph()

        # Excel name
        filename_string = f'{conversion_download_name.value}_SN{self.signal_to_noise}_WinL'
        filename_string += f'{self.signal_to_noise_win_len}_BinC{self.signal_to_noise_bin_count}_MinReqEl'
        filename_string += f'{self.signal_to_noise_min_required_elements}_SpaDif{self.spacing_difference}_Gap'
        filename_string += f'{self.spacing_difference_gap}_Miss{self.missing}.xlsx'
        document.add_paragraph(f'Peak Lists were saved as an Excel (ready for Data Alignment) named {filename_string}')

        # Save Word
        document.save(f'Files_To_Align/{conversion_download_name.value}.docx')

        # Save Excel
        with  pd.ExcelWriter('Files_To_Align/'+filename_string) as writer:
            for s in self.samples:
                self.samples[s]['centroided'].to_excel(writer, sheet_name=s)

        # Prepare the Data Alignment Step
        alignment_storage.samples = [self.samples[s]['centroided'] for s in self.samples]
        alignment_storage.update_widgets(['converted_df',])
        alignment_page.append(data_alignment_section)

        pn.state.notifications.success(f'Files successfully saved and prepared for Data Alignment.')


    def reset(self):
        "Resets all relevant parameters."
        for param in self.param:
            if param not in ["name"]:
                setattr(self, param, self.param[param].default)


    def __init__(self, **params):

        super().__init__(**params)

        widgets = {
            'signal_to_noise': pn.widgets.FloatInput(name='Minimal signal-to-noise ratio for a peak to be picked',
                start=0, value=3.3, step=0.1, styles={'font-weight': 'bold'},
                description='0 disables SNT estimation'),
            'spacing_difference_gap': pn.widgets.FloatInput(name='Spacing Difference Gap',
                start=0, value=4.0, step=0.2,
                description="""
The extension of a peak is stopped if the spacing between two subsequent data points exceeds '(Spacing Difference Gap) * min_spacing'.
'min_spacing' is the smaller of the two spacings from the peak apex to its two neighboring points.
'0' to disable the constraint."""),
            'spacing_difference': pn.widgets.FloatInput(name='Spacing Difference', start=0, value=1.5, step=0.1,
                description="""
Maximum allowed difference between points during peak extension, in multiples of the minimal difference between the peak apex and its two neighboring points.
If this difference is exceeded a missing point is assumed (see parameter 'missing'). A higher value implies a less stringent peak definition, since individual signals within the peak are allowed to be further apart.
'0' to disable the constraint."""),
            'missing': pn.widgets.IntInput(name='Missing Points Allowed', start=0, value=1, step=1,
                description="""
Maximum nº of missing points allowed when extending a peak to the left or to the right.
A missing data point occurs if the spacing between two subsequent data points exceeds '(Spacing Difference) * min_spacing'.
'min_spacing' is the smaller of the two spacings from the peak apex to its two neighboring points."""),
            'signal_to_noise_win_len': pn.widgets.FloatInput(name='Window Length in Thomson', start=1, value=200.0, step=5),
            'signal_to_noise_bin_count': pn.widgets.IntInput(name='Number of Bins for Intensity Values',
                start=3, value=30, step=1),
            'signal_to_noise_min_required_elements': pn.widgets.IntInput(
                name='Minimum Nº of Elements Required in a Window (Otherwise It is Considered Sparse)',
                start=1, value=10, step=1),
            'confirm_spectral_conv_button': pn.widgets.Button(
                name="Convert Spectral mzML data to peak lists", button_type='primary')
        }

        widgets2 = {
            'save_conversion_res_button': pn.widgets.Button(
                name="Save conversion parameters and description and the Excel with the samples' peak lists and progress to Data Alignment",
                button_type='warning', icon=download_icon)
        }

        self.controls = pn.Param(self, parameters=['signal_to_noise', 'spacing_difference_gap', 'spacing_difference',
                                                  'missing', 'signal_to_noise_win_len', 'signal_to_noise_bin_count',
                                                   'signal_to_noise_min_required_elements',
                                                   'confirm_spectral_conv_button'],
                                 widgets=widgets, name='Spectral mzML raw data conversion to peak lists Parameters')

        self.controls2 = pn.Param(self, parameters=['save_conversion_res_button'],
                                 widgets=widgets2, name='')

# Initiating mzML conversion class
mzML_store = ReadingMZML_Files()

# Click button to start conversion from mzML to mass lists
mzML_store.controls.widgets['confirm_spectral_conv_button'].on_click(mzML_store._confirm_spectral_conv_button)

# Widget to select general name of the descriptions of mzML conversion
conversion_download_name = pn.widgets.TextAreaInput(name='Name to save the excel with :',
                                                 value='mzMLconversion_PeakLists', rows=1)

mzML_store.controls2.widgets['save_conversion_res_button'].on_click(mzML_store._save_conversion_res_button)


# Contains relevant parameters regarding Data Alignment
class DataAlignmentObject(param.Parameterized):
    "Class to contain all relevant parameters regarding Data Alignment."

    # Filename
    files = param.List()

    # Samples in the list
    samples = param.List()

    # Aligned DataFrame
    aligned_df = param.DataFrame()

    # Parameters for Alignment
    ppmtol = param.Number(default=1.0) # PPM Deviation Tolerance to make the aligned buckets
    min_samples = param.Number(default=2) # Min. number of samples an aligned feature must appear to not be discarded
    confirm_button_alignment = param.Boolean(default=False)

    # Alignment Description
    desc = param.DataFrame()
    aligned_desc = param.String(default='')


    def update_widgets(self, filenames):
        "Updated min_samples widget based on number of samples in read file."

        self.files = filenames
        self.controls.widgets['min_samples'].end = len(self.samples)
        if len(self.samples) < self.min_samples:
            self.controls.widgets['min_samples'].value = len(self.samples)
            self.min_samples = len(self.samples)


    def _confirm_button_alignment(self, event):
        "Performs Data Alignment and updates layout."

        # Updating layout
        while len(alignment_page)>3:
            alignment_page.pop(-1)

        # Loading Widget while Data Alignment is happenning
        alignment_page.append(pn.indicators.LoadingSpinner(value=True, size=120,
                                                        name='Performing Data Alignment...'))

        try:
            # Perform Alignment
            aligned, desc = align(self.samples,
                          min_samples=1, 
                          ppmtol=self.ppmtol,
                          return_alignment_desc=True,
                          verbose=False)

            n_groups = len(aligned)

            # Get only featuress that appeared in at least min_samples samples
            desc = desc[desc['# features'] >= self.min_samples]
            aligned = aligned.iloc[desc.index]

            # Storing files
            self.aligned_df = aligned
            self.desc = desc

            # Building the description
            n_total_features = 0
            for samp in self.samples:
                n_total_features += len(samp)
            desc_list = [f'A total of **{n_total_features} features** were extracted in **{len(self.samples)} samples**. Data Alignment was performed with PPM Deviation tolerance of {self.ppmtol}.',
                         f'**{n_groups}** total groupings were found. From these, **{n_groups-len(self.aligned_df)}** groups were discarded (appeared in less than **{self.min_samples} samples**).',
                         f'The aligned dataset had the remaining **{len(self.aligned_df)} metabolic features**.',
                         '']
            summary_string = alignment_summary(self.desc, self.ppmtol).split('\n')
            desc_list.extend(summary_string)
            self.aligned_desc = '<br />'.join(desc_list)

            # Updating Layout
            alignment_page.pop(-1)
            alignment_page.append(self.aligned_desc)
            if len(self.aligned_df) < 8000:
                alignment_page.append(pn.pane.DataFrame(self.aligned_df, height=600))
            else:
                alignment_page.append(pn.pane.DataFrame(self.aligned_df.iloc[:8000], height=600))
            alignment_page.append(saving_df_widgets)

            if spectra_store.show_spectra:
                alignment_page.append(spectra_visualization_section)

            pn.state.notifications.success(f'Data sucessfully aligned.')

        except:
            while len(alignment_page)>2:
                alignment_page.pop(-1)
            pn.state.notifications.error(f'Data could not be aligned.')


    def reset(self):
        "Resets all relevant parameters."
        for param in self.param:
            if param not in ["name"]:
                setattr(self, param, self.param[param].default)


    def __init__(self, **params):

        super().__init__(**params)

        widgets = {
            'ppmtol': pn.widgets.FloatInput(name='PPM Deviation Tolerance (for metabolic feature grouping)',
                start=0.1, end=1000, value=1, step=0.1, styles={'font-weight': 'bold'}),
            'min_samples': pn.widgets.IntSlider(name='Minimum Sample Number Appearance (for a metabolic feature)',
                start=1, value=2, step=1, end=3, styles={'font-weight': 'bold'}),
            'confirm_button_alignment': pn.widgets.Button(name="Perform Data Alignment", button_type='warning')
        }

        self.controls = pn.Param(self, parameters=['ppmtol', 'min_samples', 'confirm_button_alignment'],
                                 widgets=widgets, name='Alignment Parameters')

# Initialize Object
alignment_storage = DataAlignmentObject()

# Data Alignment Function happens when you press the corresponding button
alignment_storage.controls.widgets['confirm_button_alignment'].on_click(alignment_storage._confirm_button_alignment)


# Needed Widgets
filename = pn.widgets.FileSelector(directory=os.getcwd()+ '/Files_To_Align', name='Select Desired Files to read')

confirm_button_filename = pn.widgets.Button(name='Read Samples', button_type='primary', disabled=False)


def _confirm_button_filename(event):
    "Reads the samples in the provided file."

    # Updating layout
    while len(alignment_page)>1:
        alignment_page.pop(-1)
    while len(data_reading_section)>2:
        data_reading_section.pop(-1)

    # Number of mzML files
    n_mzML_files = 0
    for i in filename.value:
        if i.endswith('mzML'):
            n_mzML_files += 1

    # No Files Provided
    if len(filename.value) == 0:
        pn.state.notifications.error(f'No Files were provided to perform spectral raw data conversion or data alignment.')

    # One Excel file is provided or multiple csv files (no mzML files) - Data Alignment
    elif n_mzML_files == 0:
        # A single Excel File
        if len(filename.value) == 1:
            file = filename.value[0]
            if file.endswith('.xlsx') or file.endswith('xls'):
                try:
                    # Reading the file
                    data = pd.read_excel(file, engine="openpyxl", sheet_name=None)

                    # Editing the file format a bit
                    full_data = []
                    for i in range(len(data)):
                        single_sample = list(data.values())[i]
                        single_sample = single_sample.set_index(single_sample.columns[0])
                        single_sample.columns = [list(data.keys())[i]]
                        full_data.append(single_sample)

                    # Storing the samples
                    alignment_storage.samples = full_data
                    small_name = file.split("Files_To_Align\\")[1]
                    data_reading_section.append(
                        f'**{len(full_data)} samples** were read from {small_name}.')

                    # Plotting a bar plot
                    samp_bar_plot = px.bar(x=[len(i) for i in alignment_storage.samples],
                        y=[i.columns[0] for i in alignment_storage.samples], orientation='h',
                        title='Metabolic Features per Sample')
                    samp_bar_plot.update_layout(
                        xaxis_title="Number of Metabolic Features",
                        yaxis_title="Samples",
                        height=25*len(alignment_storage.samples))
                    data_reading_section.append(pn.pane.Plotly(samp_bar_plot, config={'toImageButtonOptions': {
                            'filename': f'FeaturePerSamplePlot_{small_name}', 'scale':4}}, height=20*len(alignment_storage.samples)+200))

                    # Update the possible minimum number of samples and the layout
                    alignment_storage.update_widgets(filename.value)
                    alignment_page.append(data_alignment_section)
                    spectra_store.show_spectra = False

                    pn.state.notifications.success(f'Samples successfully read from {small_name}.')

                except:
                    small_name = file.split("Files_To_Align\\")[1]
                    pn.state.notifications.error(
                        f'Could not successfully read samples from {small_name}. Confirm formatting.')
            else:
                small_name = file.split("Files_To_Align\\")[1]
                data_reading_section.append(
                        f'**{small_name}** provided is not .xlsx or .xls file so Data Alignment cannot be performed.')
                pn.state.notifications.error(
                        f'Single File provided is not .xlsx or .xls file so Data Alignment cannot be performed.')

        # Multiple csv files
        else:
            # Number of csv files
            n_csv_files = 0
            for i in filename.value:
                if i.endswith('csv'):
                    n_csv_files += 1

            # All csv files
            if n_csv_files == len(filename.value):
                try:
                    full_data = []
                    for file in filename.value:
                        # Reading the file
                        data = pd.read_csv(file)
                        # Editing the file format a bit
                        data = data.set_index(data.columns[0])
                        full_data.append(data)

                    # Storing the samples
                    alignment_storage.samples = full_data
                    data_reading_section.append(
                        f'**{len(full_data)} samples** were read.')

                    # Plotting a bar plot
                    samp_bar_plot = px.bar(x=[len(i) for i in alignment_storage.samples],
                        y=[i.columns[0] for i in alignment_storage.samples], orientation='h',
                        title='Metabolic Features per Sample')
                    samp_bar_plot.update_layout(
                        xaxis_title="Number of Metabolic Features",
                        yaxis_title="Samples",
                        height=25*len(alignment_storage.samples))
                    data_reading_section.append(pn.pane.Plotly(samp_bar_plot, config={'toImageButtonOptions': {
                            'filename': f'FeaturePerSamplePlot_set_csv_files', 'scale':4}}, height=20*len(alignment_storage.samples)+200))

                    # Update the possible minimum number of samples and the layout
                    alignment_storage.update_widgets(filename.value)
                    alignment_page.append(data_alignment_section)
                    spectra_store.show_spectra = False

                    pn.state.notifications.success(f'Samples successfully read.')

                except:
                    pn.state.notifications.error(
                        f'Could not successfully read samples from the set of csv files. Confirm formatting.')

            # Not all csv files
            else:
                data_reading_section.append(
                    f'Multiple files were provided but they are not all sample csv files. Check Files provided.')
                pn.state.notifications.error(
                    f'Multiple files were provided but not all are csv (nor all mzML) files.')

    # Multiple files provided
    else:
        # All mzML files
        if n_mzML_files == len(filename.value):
            # Updating layout
            while len(alignment_page)>2:
                alignment_page.pop(-1)
            alignment_page.append(data_mzML_section)
            mzML_store.files = filename.value
            spectra_store.show_spectra = True
            data_reading_section.append(
                            f'**{len(filename.value)} mzML files** were detected.')
            pn.state.notifications.success(
                    f'All Files provided are mzML. Proceed to converting raw spectral data.')

        # Not all mzML files
        else:
            data_reading_section.append(
            f'Multiple files were provided but they are not all mzML spectral raw data files to convert. Check Files provided.')
            pn.state.notifications.error(
                    f'Multiple files were provided but not all are mzML (nor all csv) files.')



# Function happens when you press the button        
confirm_button_filename.on_click(_confirm_button_filename)


# Widget to select name of the aligned dataset
dataset_download_name = pn.widgets.TextAreaInput(name='Name to save the dataset as:',
                                           placeholder='aligned_df', value='aligned_df', rows=1)

# Button to save files
save_alignment_files_button = pn.widgets.Button(name='Save Aligned Dataset and Alignment Description',
                                                button_type='success', disabled=False, icon=download_icon)

def _save_alignment_files_button(event):
    "Saves aligned DataFrame and description."

    try:
        # Saving the file
        alignment_storage.aligned_df.to_csv(f'{dataset_download_name.value}.csv')
        alignment_storage.desc.to_excel(f'{dataset_download_name.value}_desc.xlsx')

        pn.state.notifications.success(f'Aligned dataset successfully saved in {dataset_download_name.value}.csv.')

    except:
        pn.state.notifications.error(f'Dataset could not be saved in {dataset_download_name.value}.csv.')

# Function happens when you press the button        
save_alignment_files_button.on_click(_save_alignment_files_button)



# Contains relevant parameters regarding Data Alignment
class Visualizing_Spectra(param.Parameterized):
    "Class to contain all relevant parameters regarding visualizing spectra."

    # Should this be shown
    show_spectra = param.Boolean(default=False)

    # Parameters for Spectra Visualization
    see_spec = param.List()
    norm_specs = param.Boolean(default=True)
    mass_lim = param.Tuple((200,1000))
    confirm_spectra_button = param.Boolean(default=False)

    # Save Parameters
    save_parameters = param.Dict()

    # Spectra
    save_fig = param.List(['To Plot the Fig'])

    def _confirm_spectra_button(self, event):
        "Plots interactive figure for spectra visualization."

        # Update Layout
        while len(spectra_visualization_section)>3:
            spectra_visualization_section.pop(-1)

        # Create figure
        fig, axs = plt.subplots(len(self.see_spec), 3, figsize=(5,3), constrained_layout=True)
        plt.rc('axes', labelsize=6, titlesize=7)
        plt.rc('xtick', labelsize=6)
        plt.rc('ytick', labelsize=6)

        # Fill in the data
        for (s, s1) in zip(self.see_spec, range(len(self.see_spec))):

            # Get data for the spectra Visualization
            spec_form = {'Raw':{}, 'Centroided':{}, 'Aligned':{}}
            spec_form['Raw']['masses'] = mzML_store.samples[s]['raw'].index
            spec_form['Centroided']['masses'] = mzML_store.samples[s]['centroided'].index
            spec_form['Aligned']['masses'] = alignment_storage.aligned_df.index
            if self.norm_specs:
                spec_form['Raw']['ints'] = mzML_store.samples[s]['raw'].values/mzML_store.samples[s]['raw'].values.sum()
                spec_form['Centroided']['ints'] = mzML_store.samples[s]['centroided'].values/mzML_store.samples[s]['raw'].values.sum()
                spec_form['Aligned']['ints'] = alignment_storage.aligned_df[s].values/mzML_store.samples[s]['raw'].values.sum()
            else:
                spec_form['Raw']['ints'] = mzML_store.samples[s]['raw'].values
                spec_form['Centroided']['ints'] = mzML_store.samples[s]['centroided'].values
                spec_form['Aligned']['ints'] = alignment_storage.aligned_df[s].values

            # Plot the Figure
            for (sf, sf1) in zip(spec_form.keys(), range(len(spec_form))):

                # If only 1 sample selected
                if len(self.see_spec) == 1:
                    axs[sf1].plot(spec_form[sf]['masses'], spec_form[sf]['ints'])
                    axs[sf1].set_xlim(self.mass_lim)
                    axs[sf1].set_title(s+' - ' + sf)
                    axs[sf1].tick_params(axis='both', which='major', labelsize=6)

                # If multiple sample selected
                else:
                    axs[s1, sf1].plot(spec_form[sf]['masses'], spec_form[sf]['ints'])
                    axs[s1, sf1].set_xlim(self.mass_lim)
                    axs[s1, sf1].set_title(s+' - ' + sf)
                    axs[s1, sf1].tick_params(axis='both', which='major', labelsize=6)

        self.save_parameters = {"see_spec":self.see_spec,
                                "norm_specs":self.norm_specs,
                                "mass_lim":self.mass_lim,}

        # Update Layout
        self.save_fig[0] = fig
        spectra_visualization_section.append(pn.pane.Matplotlib(self.save_fig[0], height=800, interactive=True, sizing_mode="stretch_width",
                                                                tight=True, dpi=300))
        spectra_visualization_section.append('')



    def update_widgets(self):
        "Updated min_samples widget based on number of samples in read file."

        self.controls.widgets['see_spec'].options = list(mzML_store.samples.keys())
        self.see_spec = list(mzML_store.samples.keys())[:3]
        self.controls.widgets['see_spec'].value = list(mzML_store.samples.keys())[:3]


    def reset(self):
        "Resets all relevant parameters."
        for param in self.param:
            if param not in ["name"]:
                setattr(self, param, self.param[param].default)


    def __init__(self, **params):

        super().__init__(**params)

        widgets = {
            'see_spec': pn.widgets.MultiChoice(name='Select Samples to See in the Spectra', value=[], options=[''],
                search_option_limit=10,
                description='Selecting too many samples (many spectra) can overload the program and lead to a crash. We recommend only selecting 3 to 4 samples.'),
            'norm_specs': pn.widgets.Checkbox(value=True,
                name='Perform Normalization to the sum of all intensities in raw data (including in centroided and aligned spectra). Used for easier spectra comparison'),
            'mass_lim': pn.widgets.IntRangeSlider(name='Lower and Upper Mass Limits to Show in the Spectra', start=0, end=2000, value=(200, 1000), step=50),
            'confirm_spectra_button': pn.widgets.Button(name="Plot the Spectra for Visualization", button_type='primary')
        }

        self.controls = pn.Param(self, parameters=['see_spec', 'norm_specs', 'mass_lim', 'confirm_spectra_button'],
                                 widgets=widgets, name='Parameters for Spectra Visualization')

# Initialize Object
spectra_store = Visualizing_Spectra()

# Data Alignment Function happens when you press the corresponding button
spectra_store.controls.widgets['confirm_spectra_button'].on_click(spectra_store._confirm_spectra_button)


# Section of the interface regarding reading the file
data_reading_section = pn.Column(filename, confirm_button_filename)

# Section of the interface regarding spectral mzML raw data conversion
data_mzML_section = pn.Column('# Converting spectral mzML raw data into mass peak lists',
                              mzML_conversion_specific_string,
                              mzML_store.controls)

# Section of the interface regarding Data Alignment
data_alignment_section = pn.Column('# Performing Data Alignment',
                                   data_alignment_specific_string, alignment_storage.controls)

# Section of the interface regarding Spectra Visualization
spectra_visualization_section = pn.Column('# Spectra Visualization',
                                   spectra_visualization_specific_string, spectra_store.controls)

# Section to save the aligned dataset
saving_df_widgets = pn.Column(dataset_download_name, save_alignment_files_button)

alignment_page = pn.Column(data_reading_section)

# Create the main area and display the first page
main_area = AlignmentPage().content

app = pn.template.BootstrapTemplate(title='Data Alignment Interface', sidebar=[], main=[main_area])

app.show(websocket_max_message_size=100*1024*1024, http_server_kwargs={'max_buffer_size': 100*1024*1024})

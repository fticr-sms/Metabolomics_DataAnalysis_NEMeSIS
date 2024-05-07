
# Necessary imports
import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from stat import S_IREAD, S_IRGRP, S_IROTH
import panel as pn
import pandas as pd
import numpy as np
import json

def ReportGenerator(folder, RepGen, file, checkbox_annotation, checkbox_formula, radiobox_neutral_mass, checkbox_others,
                     target_list, UnivarA_Store, characteristics_df, DataFrame_Store, n_databases, DB_dict, verbose_annotated_compounds,
                     data_ann_deduplicator, com_exc_compounds, PCA_params, HCA_params, PLSDA_store, RF_store, dataviz_store, PathAssign_store,
                     PCA_params_binsim, HCA_params_binsim, PLSDA_store_binsim, RF_store_binsim, rep_gen_page):
    "Makes a read-only Word file with the metabolomics data analysis performed of selected statistical analysis."

    # Create Folder to put the report in
    if not os.path.exists(folder):
        os.mkdir(folder)

    # If the report was already created in this folder
    if os.path.isfile(folder+'/Report.docx'):
        pn.state.notifications.warning((f'Report.docx already exists in {folder} folder. Please select another folder) name or erase the'
                                        f' {folder} folder from your directory.'))
        while len(rep_gen_page) > 6:
            rep_gen_page.pop(5)
        raise ValueError((f'Report.docx already exists in {folder} folder. Please select another folder) name or erase the'
                                        f' {folder} folder from your directory.'))


    # Creating the Word Document
    document = docx.Document()

    # Adding the title
    document.add_heading('Generated Data Analysis Report', 0)

    # Adding the disclaimer
    disclaimer = document.add_paragraph('This report is an attempt to describe the metabolomics data analysis performed using ')
    disclaimer.add_run('the MetsTA software. We cannot assure that the descriptions shown here are 100% correct, therefore ')
    disclaimer.add_run('revision of the report to confirm the information shown is advised. Any found inconsistencies are ')
    disclaimer.add_run('encouraged to be reported so the reason can be found and fixed.')

    # Adding the first major section
    document.add_heading('Data Pre-Processing and Pre-Treatment', level=1)

    # Chapter of section
    document.add_heading('Data Reading and Partition', level=2)

    # 1st paragraph - Dataset Read
    p1 = document.add_paragraph('Dataset analysed was present in the file: ')
    p1.add_run(RepGen.filename).bold = True
    p1.add_run('. This dataset had a total of ')

    dataset_cols = len(file.read_df.columns)
    # If a neutral mass or m/z column column was added to the dataset
    if RepGen.neutral_mass_column:
        nm_string = f' Furthermore, '
        if RepGen.type_of_mass_values_in_file == 'Neutral':
            nm_string = nm_string + '"Neutral Mass" column was added to the data '
            nm_string = nm_string + 'by directly interpreting the first column mass values given as floats (numbers).'
        elif RepGen.type_of_mass_values_in_file == 'm/z (Positive)':
            nm_string = nm_string + '"Probable m/z" column was added to the data '
            nm_string = nm_string + 'by interpreting the first column mass values given as floats (numbers) representing'
            nm_string = nm_string + ' m/z values obtained in positive ionization mode.'
        elif RepGen.type_of_mass_values_in_file == 'm/z (Negative)':
            nm_string = nm_string + '"Probable m/z" column was added to the data '
            nm_string = nm_string + 'by interpreting the first column mass values given as floats (numbers) representing'
            nm_string = nm_string + ' m/z values obtained in negative ionization mode.'
        dataset_cols = dataset_cols - 1
    else:
        nm_string = ''

    # If the target was in the first row of the dataset read or not
    if RepGen.target_included_in_file:
        dataset_size = len(file.read_df) + 1
        p1.add_run(f'{dataset_size} rows by {dataset_cols} columns including metadata and the first row (has dataset target).')
        p1.add_run(f' That is, the dataset has {dataset_size - 1} detected features.')
        
    else:
        dataset_size = len(file.read_df)
        p1.add_run(f'{dataset_size} rows by {dataset_cols} columns including metadata.')

    # Neutral / m/z mass section
    p1.add_run(nm_string)

    # 2nd paragraph - Metadata Columns
    # Putting metaadata columns in string format
    ann_cols = ', '.join(checkbox_annotation.value)
    form_cols = ', '.join(checkbox_formula.value)
    other_cols = ', '.join(checkbox_others.value)

    # Annotation and Formula columns section
    p2 = document.add_paragraph(f'From the {dataset_cols} columns, {len(checkbox_annotation.value)} (')
    p2.add_run(f'{ann_cols}').italic = True
    p2.add_run(f') was/were selected as columns with data annotations; {len(checkbox_formula.value)} (')
    p2.add_run(f'{form_cols}').italic = True
    p2.add_run(f') was/were selected as columns with formula annotations; ')

    # Neutral/m/z Mass and Other Metadata columns section
    if radiobox_neutral_mass.value != 'None':
        p2.add_run(f'{radiobox_neutral_mass.value}').italic = True
        p2.add_run(f' was selected as the column representing mass (neutral or m/z) values; and {len(checkbox_others.value)} (')
    else:
        p2.add_run(f'no column was selected to represent mass values (neutral or m/z); and {len(checkbox_others.value)} (')
    p2.add_run(f'{other_cols}').italic = True
    p2.add_run(f') was/were selected as metadata columns with other information (which were not used).')

    # 3rd paragraph - Target and Sample Columns
    class_str = ', '.join(target_list.classes)

    # Text on the target
    p3 = document.add_paragraph(f'From the {dataset_cols} columns, the remaining {len(target_list.sample_cols)} were samples.')
    p3.add_run(f'These samples belonged to {len(target_list.classes)} classes: ')
    p3.add_run(f'{class_str}').bold = True
    p3.add_run(f'. It had an average of {(len(target_list.sample_cols)/len(target_list.classes)):.2f} samples per class. ')
    p3.add_run(f'The target was the following:')

    # Presenting target as classes
    target_line = document.add_paragraph('')
    target_line.add_run('Target (Class Labels)').bold = True

    for t in range(0, len(target_list.sample_cols), 5):
        table_target = document.add_table(rows=2, cols=6, style='Table Grid')
        sample_cells = table_target.rows[0].cells
        target_cells = table_target.rows[1].cells

        sample_cells[0].text = 'Samples'
        target_cells[0].text = 'Target'
        min_r = min(6, len(target_list.sample_cols)-t+1)
        for col in range(1,min_r):
            sample_cells[col].text = str(target_list.sample_cols[col-1+t])
            target_cells[col].text = str(target_list.target[col-1+t])

        # Give space between tables
        document.add_paragraph()


    # Chapter of section
    document.add_heading('Data Filtering and Pre-Treatment', level=2)

    disclaimer2 = document.add_paragraph('')
    disclaimer2.add_run('Annotation and Annotation De-Duplication (if performed) happenned after data ').bold = True
    disclaimer2.add_run('filtering but before data pre-treatment, that is, metabolic feature merging happenned before data pre-treatment').bold = True
    disclaimer2.add_run(' despite being later in the report.').bold = True

    p4 = document.add_paragraph('Data Filtering and Pre-Treatment parameters used were as following:')

    # Data Filtering
    if UnivarA_Store.filt_method == 'None':
        document.add_paragraph('Filtering Method: None', style='List Bullet'
        )
    else:
        pre_treat_bullet = document.add_paragraph(
            f'Filtering Method: Features that appeared in at least ', style='List Bullet'
        )
        pre_treat_bullet.add_run(f'{UnivarA_Store.filt_kw} samples').bold = True
        if UnivarA_Store.filt_method == 'Total Samples':
            pre_treat_bullet.add_run(' of Total (All) Samples in the dataset')
        elif UnivarA_Store.filt_method == 'Class Samples':
            pre_treat_bullet.add_run(' of at least one Class in the dataset')

    # Missing Value Imputation
    if UnivarA_Store.mvi_method == 'Zero':
        document.add_paragraph(f'Missing Value Imputation Method: Zero Imputation', style='List Bullet'
            )
    else:
        document.add_paragraph(f'Missing Value Imputation Method: {UnivarA_Store.mvi_kw*100}% of {UnivarA_Store.mvi_method}',
                            style='List Bullet')

    # Normalization Method
    norm_paragraph = document.add_paragraph(f'Normalization Method:', style='List Bullet'
            )
    if UnivarA_Store.norm_method == 'Reference Feature':
        norm_paragraph.add_run(f' Normalization by a Reference Feature - {UnivarA_Store.norm_kw}')
    elif UnivarA_Store.norm_method == 'Total Intensity Sum':
        norm_paragraph.add_run(f' Normalization by Total Sum of Intensities')
    elif UnivarA_Store.norm_method == 'PQN':
        if UnivarA_Store.norm_kw in ['mean', 'median']:
            n_kw = 'the ' + str(UnivarA_Store.norm_kw[0]) + ' of all samples as reference'
        else:
            n_kw = str(UnivarA_Store.norm_kw[0]) + ' sample as reference'
        norm_paragraph.add_run(f' PQN (Probabilistic Quotient Normalization) using {n_kw}')
    elif UnivarA_Store.norm_method == 'Quantile':
        norm_paragraph.add_run(f' Quantile Normalization using the {str(UnivarA_Store.norm_kw[0])} of all samples as reference')
    elif UnivarA_Store.norm_method == 'None':
        norm_paragraph.add_run(f' None')
    else:
        pn.state.notifications.error('Normalization method used not recognized.')
        raise ValueError('Normalization method used not recognized for Report Generation.')

    # Transformation Method
    if UnivarA_Store.tf_method == 'Generalized Logarithmic Transformation (glog)':
        if UnivarA_Store.tf_kw:
            document.add_paragraph(
                f'Transformation Method: Generalized Logarithmic Transformation (glog) with lambda equal to {UnivarA_Store.tf_kw}',
                                style='List Bullet'
                )
        else:
            document.add_paragraph(
                f'Transformation Method: Generalized Logarithmic Transformation (glog) with lambda equal to 0',
                                style='List Bullet'
                )
    elif not UnivarA_Store.tf_method:
        document.add_paragraph(f'Transformation Method: None', style='List Bullet'
                )
    else:
        pn.state.notifications.error('Transformation method used not recognized.')
        raise ValueError('Transformation method used not recognized for Report Generation.')

    # Scaling Method
    if UnivarA_Store.scaling_method == 'Level Scaling':
        document.add_paragraph(
                f"Scaling Method: Level Scaling (using feature's {UnivarA_Store.scaling_kw} as scaling factors)",
                                style='List Bullet'
                )
    else:
        document.add_paragraph(f"Scaling Method: {UnivarA_Store.scaling_method}", style='List Bullet'
                )

    # Dataset extra characteristics
    range_min, range_max = characteristics_df.value.loc['feature value ranges'][0].split(' - ')
    p5 = document.add_paragraph(f"After pre-treatment, the dataset had {len(DataFrame_Store.treated_df.columns)} features. ")
    p5.add_run(f'Absolute intensity values (without pre-treatment) ranged between {range_min[1:]} to {range_max[:-1]}, ')
    p5.add_run(f'averaging at {characteristics_df.value.loc["feature value average (std)"][0]}.')
        
    document.add_paragraph(f"Treated DataFrame was downloaded as 'Report_treated_df.csv'.")


    # Chapter of section
    document.add_heading('Data Annotation and De-Duplication', level=2)

    ann_pg = document.add_paragraph(f'Data Annotation in this software was made using {n_databases.value} database(s).')

    if n_databases.value != 0:
        # Database Annotation based on PPM Deviation
        if RepGen.annotation_margin_method == 'PPM Deviation':
            ann_pg.add_run(f' Database matching was performed using a maximum of {RepGen.annotation_margin_ppm_deviation} ppm')
            ann_pg.add_run(f' (parts per million) deviation of the theoretical compound mass (based on formula) to the ')
            ann_pg.add_run(f'experimental mass in the Mass Column ({radiobox_neutral_mass.value}) chosen.')

        # Database Annotation based on flat Da deviation
        else:
            ann_pg.add_run(f' Database matching was performed using a maximum of {RepGen.annotation_margin_ppm_deviation} Da')
            ann_pg.add_run(f' deviation of the theoretical compound mass (based on formula) to the ')
            ann_pg.add_run(f'experimental mass in the Mass Column ({radiobox_neutral_mass.value}) chosen.')

    # Adducts considered to be searched
    n_adducts = len(RepGen.adducts_to_consider)
    if n_adducts != 0:
        ann_pg.add_run(f' Matching was performed against the databases considering {n_adducts} possible adducts. These were: ')
        a = 0
        for ad, mass_shift in RepGen.adducts_to_consider.items():
            ann_pg.add_run(f'{ad} (mass shift of {float(mass_shift)})')
            if a == n_adducts - 1:
                ann_pg.add_run(f'.')
            elif a == n_adducts - 2:
                ann_pg.add_run(f' and ')
            else:
                ann_pg.add_run(f', ')
            a += 1


    # Information on Databases used
    for db in range(n_databases.value):
        key = str(db + 1)
        db_info = DB_dict[key]
        document.add_paragraph(f"{db_info.abv}: Database read from '{db_info.file}' using {db_info.IDcol} column as IDs. It had {len(db_info.db_len)} metabolites.",
                            style='List Bullet')

    # Information on the annotations made with each Database used
    if n_databases.value != 0:
        document.add_paragraph(f'Annotation with these databases led to the annotation of :')
        for db in range(n_databases.value):
            db_info = DB_dict[str(db + 1)]
            document.add_paragraph(f'{db_info.abv}: {verbose_annotated_compounds[db+1].value.split( )[1]} compounds',
                                style='List Bullet')

    # Annotation De-Duplication section
    p6 = document.add_paragraph(f'After the annotations performed, the annotations (including previous annotations made) were ')
    p6.add_run('checked. The following table was obtained from observing the multiplicity of the annotations.')

    # Table with the report of the multiple annotations
    table_ann = document.add_table(rows=4, cols=len(data_ann_deduplicator.mult_ann_report.columns)+1, style='Light Grid')

    def fill_word_table(table, df):
        "Function to convert a pandas DataFrame to a table in the Word Report."
        
        # Get the Index column and fill it
        hdr_cells = table.columns[0].cells
        idx_pos = 1
        for idx in df.index:
            hdr_cells[idx_pos].text = idx
            idx_pos += 1

        # Go through every other column and fill it
        n_col = 1
        for col in df.columns:
            # Get the col
            new_col = table.columns[n_col].cells
            n_col += 1
            new_col[0].text = col # Name of the column
            a = 1
            # Fill remaining values
            for value in df[col]:
                new_col[a].text = str(value)
                a += 1
        table.autofit = True
        return table

    table_ann = fill_word_table(table_ann, data_ann_deduplicator.mult_ann_report)

    # Give space after table
    document.add_paragraph()

    # Annotation De-duplication section
    if RepGen.deduplication_performed == 'Annotation De-Duplication was not performed':
        document.add_paragraph(f'Multiple Annotation De-Duplication and Metabolic Feature merging was not performed.')

    else:
        # If annotation De-Duplication was made

        # Type of metabolic merging performed
        if RepGen.deduplication_performed == 'Annotation De-Duplication was performed but merge problems were not individually observed/decided':
            p7 = document.add_paragraph(f'Multiple Annotation De-Duplication and Metabolic Feature merging was performed but')
            p7.add_run(' possible merge problems were not individually observed. ')
        else:
            p7 = document.add_paragraph(f'Multiple Annotation De-Duplication and Metabolic Feature merging was performed and')
            p7.add_run(' merge problems were individually observed and decided whether they were merged. ')

        # If Formula columns were considered
        if data_ann_deduplicator.current_params['consider_formula_cols']:
            p7.add_run('Previously Formula annotated columns were also considered in metabolic feature merging. ')
        else:
            p7.add_run('Previously Formula annotated columns were not considered in metabolic feature merging. ')

        # If Scenario 1 of Case 4 of mergings was considered as potential problem
        if data_ann_deduplicator.current_params['problem_condition'] == 'Scenario 1 of Situation 4 like cases are not merged and are not shown.':
            p7.add_run('Furthermore, Scenario 1 of Situation 4 like cases (see description in MetsTA software for further ')
            p7.add_run('details) were not merged and are not shown (not considered possible problems).')
        else:
            p7.add_run('Furthermore, Scenario 1 of Situation 4 like cases (see description in MetsTA software for further ')
            p7.add_run('details) are shown as possible problems to individually decide.')

        # Describing Mergings
        target_line = document.add_paragraph('')
        target_line.add_run('Describing Metabolic Feature Mergings').bold = True

        # Table with descriptions of types of merging made
        table_merge_situations = document.add_table(rows=len(data_ann_deduplicator.merge_situations.index)+1,
                                                    cols=len(data_ann_deduplicator.merge_situations.columns)+1,
                                                    style='Light Grid')
        table_merge_situations = fill_word_table(table_merge_situations, data_ann_deduplicator.merge_situations)

        # Give space after table
        document.add_paragraph()

        # Merge Report - Treating the string first
        m_report = data_ann_deduplicator.merge_report
        m_report = m_report.replace('**', '')
        m_report = m_report.replace('  ', '')
        document.add_paragraph(m_report)

        # Saving merge description as Excel
        data_ann_deduplicator.merge_description.to_excel(folder+'/Report_Metabolic_Feature_Peak_Merging_Description.xlsx')
        document.add_paragraph("Detailed information of metabolic feature peak merging made saved in 'Report_Metabolic_Feature_Peak_Merging_Description.xlsx'.")


    # Final Description
    p8 = document.add_paragraph("")
    p8.add_run("Finished Data Pre-Treatment").bold = True
    p8.add_run(f". Dataset had {len(target_list.sample_cols)} samples and {len(DataFrame_Store.treated_df.columns)} metabolic features,")
    p8.add_run(f" from which {sum(DataFrame_Store.metadata_df['Has Match?'])} are annotated.")
    p8.add_run(f" Dataset after annotation and metabolic feature merging (if performed) were saved. 'Report_Complete_non_Treated_df.csv'")
    p8.add_run(f" has original intensity values and metadata. 'Report_Complete_Treated_df.csv' has treated intensity values and ")
    p8.add_run(f"metadata. 'Report_Treated_df.csv' only has the treated intensity values.")

    # Saving the respective DataFrames
    DataFrame_Store.original_df.to_csv(folder+'/Report_Complete_non_Treated_df.csv')
    DataFrame_Store.treated_df.to_csv(folder+'/Report_Treated_df.csv')
    pd.concat((DataFrame_Store.metadata_df, DataFrame_Store.treated_df.T), axis=1).to_csv(folder+'/Report_Complete_Treated_df.csv')

    # End of section
    document.add_page_break()




    # Adding the second major section
    document.add_heading('Statistical Analysis', level=1)

    # Getting the Statistical Methods that will be in 
    stat_methods = []
    for i in RepGen.com_exc_analysis:
        if i == 'Overview':
            stat_methods.append('Common and Exclusive Compound Overview')
        else:
            stat_methods.append(i)
    for i in RepGen.unsup_analysis:
        stat_methods.append(i)
    for i in RepGen.sup_analysis:
        stat_methods.append(i)
    if RepGen.univ_analysis:
        stat_methods.append('Univariate Analysis')
    for i in RepGen.dataviz_analysis:
        stat_methods.append(i)
    if RepGen.pathassign_analysis:
        stat_methods.append('HMDB Pathway Assignment')
    for i in RepGen.BinSim_analysis:
        stat_methods.append('BinSim - ' + i)

    document.add_paragraph(f'This section includes the statistical analysis chosen to be in this report. These are: {", ".join(stat_methods)}.')

    # Common and Exclusive Compound Section
    if len(np.intersect1d(['Common and Exclusive Compound Overview', 'Venn Diagram', 'Intersection Plot'], stat_methods)) != 0:
        document.add_heading('Common and Exclusive Compound Analysis', level=2)

        # Overview Section
        if 'Common and Exclusive Compound Overview' in stat_methods:
            # Heading
            document.add_heading('Common and Exclusive Compound Overview', level=3)

            # In case analysis was not performed
            if len(com_exc_compounds.com_exc_desc) == 0:
                document.add_paragraph('Common and Exclusive Compound Analysis was not performed. Thus, this section will be skipped.')

            else:
                # Descriptions of peaks per class and exclusive peaks per class
                document.add_paragraph(com_exc_compounds.groups_description.replace('<br />', '\n').replace('**', ''))
                document.add_paragraph()
                document.add_paragraph(com_exc_compounds.com_exc_desc.replace('<br />', '\n').replace('**', ''))

        # Venn Diagram Section
        if 'Venn Diagram' in stat_methods:
            # Heading
            document.add_heading('Venn Diagram', level=3)

            # If there is a Venn Diagram
            if type(com_exc_compounds.Venn_plot[0]) != str:
                # Description
                venn_pg = document.add_paragraph('Venn Diagram downloaded as currently present in the program. Venn Diagram ')
                venn_pg.add_run(f"shown includes the following classes: ")
                venn_pg.add_run(f"{', '.join(com_exc_compounds.venn_class_subset)}.").bold = True

                # Creating filename for the Venn Figure and saving it
                filename_string = f'/Report_Venn_diagram_{com_exc_compounds.type_of_venn}_classes'
                for cl in com_exc_compounds.venn_class_subset:
                    filename_string = filename_string + '_'+cl
                com_exc_compounds.Venn_plot[0].savefig(folder+filename_string, dpi=com_exc_compounds.dpi_venn)

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(15))

            # If there is no Venn Diagram
            else:
                venn_pg = document.add_paragraph('Venn Diagram was not created during the Data Analysis. ')
                venn_pg.add_run(f"Thus, this section will be skipped.")

        # Intersection Plot Section
        if 'Intersection Plot' in stat_methods:
            # Heading
            document.add_heading('Intersection Plots', level=3)

            # If there are Intersection Plots
            if type(com_exc_compounds.IntersectionPlot[0]) != str:
                # Description
                inter_pg = document.add_paragraph('Intersection Plots downloaded as currently present in the program. ')
                inter_pg.add_run(f"Intersection Plots shown include the following classes: ")
                inter_pg.add_run(f"{', '.join(com_exc_compounds.inter_class_subset)}.").bold = True

                # Creating filename for the Intersection Plots and saving them
                filename_string_all_mets = f'/Report_IntersectionPlot_all_metabolites_classes'
                filename_string_ann_mets = f'/Report_IntersectionPlot_annotated_metabolites_classes'
                for cl in com_exc_compounds.inter_class_subset:
                    filename_string_all_mets = filename_string_all_mets + '_'+cl
                    filename_string_ann_mets = filename_string_ann_mets + '_'+cl
                com_exc_compounds.IntersectionPlot[0].savefig(folder + filename_string_all_mets + '.png',
                                                            dpi=com_exc_compounds.dpi_inter)
                if type(com_exc_compounds.IntersectionPlot[1]) != str:
                    com_exc_compounds.IntersectionPlot[1].savefig(folder + filename_string_ann_mets + '.png',
                                                                dpi=com_exc_compounds.dpi_inter)

                # Adding description and figure - all metabolites
                inter_pg2 = document.add_paragraph()
                inter_pg2.add_run('Intersection Plot with all metabolites of the dataset').bold = True
                inter_pg2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_picture(folder+filename_string_all_mets + '.png', width=Cm(14))

                # Adding description and figure - annotated metabolites
                if type(com_exc_compounds.IntersectionPlot[1]) != str:
                    inter_pg3 = document.add_paragraph()
                    inter_pg3.add_run('Intersection Plot with only annotated metabolites in the dataset').bold = True
                    inter_pg3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    document.add_picture(folder+filename_string_ann_mets + '.png', width=Cm(14))
                else:
                    inter_pg3 = document.add_paragraph('Intersection Plot with only annotated metabolites in the dataset')
                    inter_pg3.add_run(' was not created since annotations were not found in the dataset.')

            # If there is no Intersection Plots
            else:
                inter_pg = document.add_paragraph('Intersection Plots were not created during the Data Analysis. ')
                inter_pg.add_run(f"Thus, this section will be skipped.")

        # End of section
        document.add_page_break()


    # Unsupervised Analysis Section
    if len(np.intersect1d(['PCA', 'HCA'], stat_methods)) != 0:
        document.add_heading('Unsupervised Analysis', level=2)

        # PCA Section
        if 'PCA' in stat_methods:
            # Heading
            document.add_heading('Principal Component Analysis (PCA)', level=3)

            # If there is a PCA Projection
            if type(PCA_params.PCA_plot[0]) != str:
                # Description
                pca_pg = document.add_paragraph(f'Principal Component Analysis was made with ')
                pca_pg.add_run(f'{len(PCA_params.controls.widgets["PCx"].options)} components.')

                pca_pg2 = document.add_paragraph()
                if PCA_params.n_dimensions == '2 Components':
                    pca_pg2.add_run(f'2-D PCA Projection Plot of Components: {PCA_params.PCx} and {PCA_params.PCy}').bold = True
                else:
                    pca_pg2.add_run(f'3-D PCA Projection Plot of Components: {PCA_params.PCx}, {PCA_params.PCy} and {PCA_params.PCz}').bold = True
                pca_pg2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Creating filename for the PCA Projection and saving it
                filename_string = '/Report_PCA_plot'
                if PCA_params.ellipse_draw:
                    if PCA_params.confidence != 0:
                        filename_string = filename_string + f'_ellipse({PCA_params.confidence*100}%confidence)'
                    else:
                        filename_string = filename_string + f'_ellipse({PCA_params.confidence_std}std)'
                PCA_params.PCA_plot[0].write_image(folder+filename_string+'.png', scale=4)
                PCA_params.PCA_plot[0].write_html(folder+filename_string+".html")

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(12.5))

                # If there is an explained variance plot
                if type(PCA_params.PCA_plot[0]) != str:
                    # Description
                    pca_pg3 = document.add_paragraph()
                    pca_pg3.add_run('Cumulative Explained Variance (by Principal Component) Plot').bold = True
                    pca_pg3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Creating filename for the PCA explained variance plot and saving it
                    filename_string = '/Report_PCA_exp_var_plot'
                    PCA_params.exp_var_fig_plot[0].write_image(folder+filename_string+'.png', scale=4)
                    PCA_params.exp_var_fig_plot[0].write_html(folder+filename_string+".html")

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(12.5))

                # If there is a PCA Projection (Scatter) Plot of the Principal Components
                if type(PCA_params.scatter_PCA_plot[0]) != str:
                    # Description
                    pca_pg4 = document.add_paragraph()
                    pca_pg4.add_run('2-D PCA Projection (Scatter) Plot of the Principal Components').bold = True
                    pca_pg4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Creating filename for the PCA Projection (Scatter) Plot of the Principal Components and saving it
                    filename_string = '/Report_PCA_scatter_plot'
                    PCA_params.scatter_PCA_plot[0].write_image(folder+filename_string+'.png', scale=4)
                    PCA_params.scatter_PCA_plot[0].write_html(folder+filename_string+".html")

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(15))

            # If there is no PCA analysis performed
            else:
                pca_pg = document.add_paragraph('Principal Component Analysis was not performed during the Data Analysis. ')
                pca_pg.add_run(f"Thus, this section will be skipped.")


        # HCA Section
        if 'HCA' in stat_methods:
            # Heading
            document.add_heading('Hierarchical Clustering Analysis (HCA)', level=3)

            # If there is a dendrogram
            if type(HCA_params.HCA_plot[0]) != str:
                # Description
                hca_pg = document.add_paragraph(f'Hierarchical Clustering Analysis was made with ')
                hca_pg.add_run(f'{HCA_params.dist_metric} distance metric and {HCA_params.link_metric} linkage metric.')

                # Creating filename for the dendrogram and saving it
                filename_string = f'/Report_HCA_plot_{HCA_params.dist_metric}Dist_{HCA_params.link_metric}Linkage'
                HCA_params.HCA_plot[0].savefig(folder + filename_string + '.png', dpi=HCA_params.dpi)

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(12))

        # End of section
        document.add_page_break()



    # Supervised Analysis Section
    if len(np.intersect1d(['PLS-DA', 'Random Forest'], stat_methods)) != 0:
        document.add_heading('Supervised Analysis', level=2)

        # PLS-DA Section
        if 'PLS-DA' in stat_methods:
            # Heading
            document.add_heading('Partial Least Squares - Discriminant Analysis (PLS-DA)', level=3)

            # If there is a PLS optimization
            if type(PLSDA_store.optim_figure[0]) != str:
                # Minor Heading
                document.add_heading('Number of Components Optimization', level=4)

                # Description
                n_min_components, n_max_components = PLSDA_store.current_other_plsda_params["n_min_max_components"]
                plsda_pg = document.add_paragraph(f'First, an optimization of the number of components to use to fit a ')
                plsda_pg.add_run(f'PLS-DA model was made by fitting PLS models from {n_min_components}')
                plsda_pg.add_run(f' to {n_max_components} components and evaluating their Q2 (mainly) and R2 scores estimated')
                plsda_pg.add_run(f' by {PLSDA_store.current_other_plsda_params["n_fold_optim"]}-fold stratified')
                plsda_pg.add_run(f' cross-validation (scale = {PLSDA_store.current_other_plsda_params["scale_optim"]}). Results')
                plsda_pg.add_run(f' are shown in the figure below with a maximum Q2 value with {PLSDA_store.rec_components} ')
                plsda_pg.add_run(f'components.')

                # Plot title
                plsda_pg2 = document.add_paragraph()
                plsda_pg2.add_run('PLS-DA Component Optimization Plot').bold = True
                plsda_pg2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Creating filename for the PLS optimization plot and saving it
                filename_string = f'/Report_PLS_optim_plot_{n_min_components}to{n_max_components}'
                filename_string = filename_string + f'components_{PLSDA_store.current_other_plsda_params["n_fold_optim"]}-'
                filename_string = filename_string + f'foldstratCV_scale{PLSDA_store.current_other_plsda_params["scale_optim"]}'
                PLSDA_store.optim_figure[0].write_image(folder+filename_string+'.png', scale=4)
                PLSDA_store.optim_figure[0].write_html(folder+filename_string+".html")

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(15))

            else:
                document.add_paragraph(f'Optimization of the number of components to fit the PLS-DA model was not performed.')


            # If the PLS-DA model was fitted
            if type(PLSDA_store.models[0]) != str:
                # Minor Heading
                document.add_heading('PLS-DA Model Fitting and Results', level=4)

                # Description
                plsda_params = PLSDA_store.current_plsda_params
                plsda_pg3 = document.add_paragraph(f'PLS-DA model was fitted using {plsda_params["n_components"]} components')
                plsda_pg3.add_run(f' estimating model performance and feature importance ({plsda_params["feat_imp"]}')
                plsda_pg3.add_run(f' score method) by {plsda_params["n_folds"]}-fold stratified cross-validation (scale = ')
                plsda_pg3.add_run(f'{plsda_params["scale"]}) repeated {plsda_params["n_iterations"]} times (randomized folds ')
                plsda_pg3.add_run(f'in cross-validation). Model Performance results are shown in the table below (by the ')

                # Building the filename and saving PLS-DA Feature Importance Table
                filename_string = f'/Report_PLS-DA_FeatImp_{plsda_params["feat_imp"]}_model_params_components'
                filename_string = filename_string + f'{plsda_params["n_components"]}_{plsda_params["n_folds"]}-foldstratCV_'
                filename_string = filename_string + f'iterations{plsda_params["n_iterations"]}_scale{plsda_params["scale"]}.xlsx'
                PLSDA_store.feat_impor.to_excel(folder+filename_string)

                # Finishing up the description
                plsda_pg3.add_run(f"metrics chosen) and Feature Importance Table was saved as '{filename_string[1:]}'. ")
                plsda_pg3.add_run(f'Furthermore, below we also show a PLS Projection of chosen Latent Variables.')

                # Table with results of PLS-DA
                table_plsda_results = document.add_table(rows=len(PLSDA_store.n_results.index)+1,
                                                            cols=len(PLSDA_store.n_results.columns)+1,
                                                            style='Light Grid')
                table_plsda_results = fill_word_table(table_plsda_results, PLSDA_store.n_results)

                # Give space after table
                document.add_paragraph()

                plsda_pg4 = document.add_paragraph()
                if PLSDA_store.n_dimensions == '2 Components':
                    plsda_pg4.add_run(f'2-D PLS Projection Plot of Latent Variables: {PLSDA_store.LVx} and {PLSDA_store.LVy}').bold = True
                else:
                    plsda_pg4.add_run(f'3-D PCA Projection Plot of Latent Variables: {PLSDA_store.LVx}, {PLSDA_store.LVy}').bold = True
                    plsda_pg4.add_run(f' and {PLSDA_store.LVz}').bold = True
                plsda_pg4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Creating filename for the PLS Projection and saving it
                filename_string = '/Report_PLS_projection_plot'
                if PLSDA_store.n_dimensions == '2 Components':
                    if PLSDA_store.ellipse_draw:
                        if PLSDA_store.confidence != 0:
                            filename_string = filename_string + f'_ellipse({PLSDA_store.confidence*100}%confidence)'
                        else:
                            filename_string = filename_string + f'_ellipse({PLSDA_store.confidence_std}std)'
                PLSDA_store.PLS_plot[0].write_image(folder+filename_string+'.png', scale=4)
                PLSDA_store.PLS_plot[0].write_html(folder+filename_string+".html")

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(15))


                # If there is Permutation Test plot
                if type(PLSDA_store.perm_figure[0]) != str:
                    # Minor Heading
                    document.add_heading('PLS-DA Permutation Test', level=4)

                    # Description
                    plsda_pg5 = document.add_paragraph('Permutation Test was performed with the same number of components as before - ')
                    plsda_pg5.add_run(f"{PLSDA_store.current_plsda_params_permutation['n_components']} - and with ")
                    plsda_pg5.add_run(f"{PLSDA_store.current_plsda_params_permutation['n_permutations']} permutations. Model ")
                    plsda_pg5.add_run(f"performance was evaluated by {PLSDA_store.current_plsda_params_permutation['perm_metric']}")
                    plsda_pg5.add_run(f" estimated with {PLSDA_store.current_plsda_params_permutation['n_folds']}-fold stratified ")
                    plsda_pg5.add_run(f"cross-validation (scale = {PLSDA_store.current_plsda_params_permutation['scale']}).")

                    # Creating filename for the PLS-DA Permutation Test and saving it
                    filename_string = f'/Report_PLS-DA_permutation_test_{PLSDA_store.current_plsda_params_permutation["n_permutations"]}perm_'
                    filename_string = filename_string + f'{PLSDA_store.current_plsda_params_permutation["n_components"]}comp_'
                    filename_string = filename_string + f'{PLSDA_store.current_plsda_params_permutation["n_folds"]}-foldstratCV_scale'
                    filename_string = filename_string + f'{PLSDA_store.current_plsda_params_permutation["scale"]}_metric'
                    filename_string = filename_string + f'{PLSDA_store.current_plsda_params_permutation["perm_metric"]}'
                    PLSDA_store.perm_figure[0].savefig(folder+filename_string+'.png', dpi=PLSDA_store.dpi)

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(15))


                # If there is ROC Curve plot
                if type(PLSDA_store.ROC_figure[0]) != str:
                    # Minor Heading
                    document.add_heading('PLS-DA ROC Curve', level=4)

                    # Description
                    roc_params = PLSDA_store.current_other_plsda_params['ROC_filename'].split('_')
                    if len(target_list.classes) > 2:
                        plsda_pg6 = document.add_paragraph('Since there are more than 2 classes, ROC curves were computed with a ')
                        plsda_pg6.add_run(f"1vsAll scheme for each of the {len(target_list.classes)} classes.")
                    else:
                        plsda_pg6 = document.add_paragraph('Since there are only 2 classes, ROC curves were computed considering ')
                        plsda_pg6.add_run(f"{roc_params[2][:-8]}").bold = True
                        plsda_pg6.add_run(f" as the positive class.")
                    plsda_pg6.add_run(f' Other parameters were maintained: number of components - {roc_params[3][:-10]}; ')
                    plsda_pg6.add_run(f"{roc_params[4][:-7]} stratified cross-validation; scale - {roc_params[6][5:]}. ")
                    plsda_pg6.add_run(f"Finally, ROC Curve estimation was repeated {roc_params[5][:-10]} times (randomized ")
                    plsda_pg6.add_run(f"cross-validation folds).")

                    # Creating filename for the PLS-DA ROC Curve and saving it
                    filename_string = '/Report_' + PLSDA_store.current_other_plsda_params['ROC_filename']
                    PLSDA_store.ROC_figure[0].write_image(folder+filename_string+'.png', scale=4)
                    PLSDA_store.ROC_figure[0].write_html(folder+filename_string+".html")

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(15))

            # If there is no PLS-DA model fitted
            else:
                plsda_pg3 = document.add_paragraph('PLS-DA model was not fitted during the Data Analysis. ')
                plsda_pg3.add_run(f"Thus, this section will be skipped.")


        # Random Forest Section
        if 'Random Forest' in stat_methods:
            # Heading
            document.add_heading('Random Forest (RF)', level=3)

            # If there is a Random Forest Tree Number optimization
            if type(RF_store.optim_figure[0]) != str:
                # Minor Heading
                document.add_heading('Number of Trees Optimization', level=4)

                # Description
                n_min_trees, n_max_trees = RF_store.current_other_rf_params["n_min_max_trees"]
                rf_pg = document.add_paragraph(f'First, an optimization of the number of trees to use to fit a ')
                rf_pg.add_run(f'Random Forest model was made by fitting RF models from {n_min_trees} to {n_max_trees} trees in ')
                rf_pg.add_run(f'{RF_store.current_other_rf_params["n_interval"]} tree steps, evaluating their model accuracy')
                rf_pg.add_run(f' estimated by {RF_store.current_other_rf_params["n_fold_optim"]}-fold stratified')
                rf_pg.add_run(f' cross-validation. Results are shown in the figure below.')

                # Plot title
                rf_pg2 = document.add_paragraph()
                rf_pg2.add_run('Random Forest Number of Trees Optimization Plot').bold = True
                rf_pg2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Creating filename for the Random Forest optimization plot and saving it
                filename_string = f'/Report_RF_optim_plot_{RF_store.current_other_rf_params["n_fold_optim"]}-foldStratCV_'
                filename_string = filename_string + f'{n_min_trees}to{n_max_trees}trees_'
                filename_string = filename_string + f'({RF_store.current_other_rf_params["n_interval"]}interval)'
                RF_store.optim_figure[0].write_image(folder+filename_string+'.png', scale=4)
                RF_store.optim_figure[0].write_html(folder+filename_string+".html")

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(15))

            else:
                document.add_paragraph(f'Optimization of the number of trees to fit the Random Forest model was not performed.')


            # If the Random Forest model was fitted
            if type(RF_store.models[0]) != str:
                # Minor Heading
                document.add_heading('Random Forest Model Fitting and Results', level=4)

                # Description
                rf_params = RF_store.current_rf_params
                rf_pg3 = document.add_paragraph(f'Random Forest model was fitted with {rf_params["n_trees"]}')
                rf_pg3.add_run(f' trees estimating model performance and feature importance by {rf_params["n_folds"]}')
                rf_pg3.add_run(f'-fold stratified cross-validation repeated {rf_params["n_iterations"]} times (randomized ')
                rf_pg3.add_run(f'folds in cross-validation). Model Performance results are shown in the table below (by the ')

                # Building the filename and saving Random Forest Feature Importance Table
                filename_string = f'/Report_RF_FeatImp_Gini_model_params_{rf_params["n_trees"]}trees_{rf_params["n_folds"]}'
                filename_string = filename_string + f'-foldstratCV_iterations{rf_params["n_iterations"]}.xlsx'
                RF_store.feat_impor.to_excel(folder+filename_string)

                # Finishing up the description
                rf_pg3.add_run(f"metrics chosen) and Feature Importance Table was saved as '{filename_string[1:]}'.")

                # Table with results of Random Forest
                table_rf_results = document.add_table(rows=len(RF_store.n_results.index)+1,
                                                            cols=len(RF_store.n_results.columns)+1,
                                                            style='Light Grid')
                table_rf_results = fill_word_table(table_rf_results, RF_store.n_results)

                # Give space after table
                document.add_paragraph()


                # If there is Permutation Test plot
                if type(RF_store.perm_figure[0]) != str:
                    # Minor Heading
                    document.add_heading('Random Forest Permutation Test', level=4)

                    # Description
                    rf_params_perm = RF_store.current_rf_params_permutation
                    rf_pg4 = document.add_paragraph('Permutation Test was performed with the same number of trees as before - ')
                    rf_pg4.add_run(f"{rf_params_perm['n_trees']} - and with {rf_params_perm['n_permutations']} permutations. ")
                    rf_pg4.add_run(f"Model performance was evaluated by {rf_params_perm['perm_metric']} estimated with ")
                    rf_pg4.add_run(f"{rf_params_perm['n_folds']}-fold stratified cross-validation.")

                    # Creating filename for the Random Forest Permutation Test and saving it
                    filename_string = f'/Report_RF_permutation_test_{rf_params_perm["n_permutations"]}perm_'
                    filename_string = filename_string + f'{rf_params_perm["n_trees"]}comp_'
                    filename_string = filename_string + f'{rf_params_perm["n_folds"]}-foldstratCV_metric'
                    filename_string = filename_string + f'{rf_params_perm["perm_metric"]}'
                    RF_store.perm_figure[0].savefig(folder+filename_string+'.png', dpi=RF_store.dpi)

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(15))


                # If there is ROC Curve plot
                if type(RF_store.ROC_figure[0]) != str:
                    # Minor Heading
                    document.add_heading('Random Forest ROC Curve', level=4)

                    # Description
                    roc_params = RF_store.current_other_rf_params['ROC_filename'].split('_')
                    if len(target_list.classes) > 2:
                        rf_pg5 = document.add_paragraph('Since there are more than 2 classes, ROC curves were computed with a ')
                        rf_pg5.add_run(f"1vsAll scheme for each of the {len(target_list.classes)} classes.")
                    else:
                        rf_pg5 = document.add_paragraph('Since there are only 2 classes, ROC curves were computed considering ')
                        rf_pg5.add_run(f"{roc_params[2][:-8]}").bold = True
                        rf_pg5.add_run(f" as the positive class.")
                    rf_pg5.add_run(f' Other parameters were maintained: number of trees - {roc_params[3][:-5]}; ')
                    rf_pg5.add_run(f"{roc_params[4][:-7]} stratified cross-validation. Finally, ROC Curve estimation was ")
                    rf_pg5.add_run(f"repeated {roc_params[5][:-10]} times (randomized cross-validation folds).")

                    # Creating filename for the Random Forest ROC Curve and saving it
                    filename_string = '/Report_' + RF_store.current_other_rf_params['ROC_filename']
                    RF_store.ROC_figure[0].write_image(folder+filename_string+'.png', scale=4)
                    RF_store.ROC_figure[0].write_html(folder+filename_string+".html")

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(15))

            # If there is no Random Forest model fitted
            else:
                rf_pg3 = document.add_paragraph('Random Forest model was not fitted during the Data Analysis. ')
                rf_pg3.add_run(f"Thus, this section will be skipped.")

        # End of section
        document.add_page_break()



    # Univariate Analysis Section
    if 'Univariate Analysis' in stat_methods:
        # Heading
        document.add_heading('Univariate Analysis', level=2)

        # In case analysis was not performed
        if len(UnivarA_Store.current_univ_params) == 0:
            document.add_paragraph('Univariate Analysis was not performed. Thus, this section will be skipped.')

        else:
            # Descriptions of Univariate Test Performed
            univ_parameters = UnivarA_Store.current_univ_params
            univ_pg = document.add_paragraph('Univariate Analysis was performed comparing the ')
            univ_pg.add_run(f'Test Class: {univ_parameters["Test Class"]}').bold = True
            univ_pg.add_run(f' against the ')
            univ_pg.add_run(f'Control Class: {univ_parameters["Control Class"]}.').bold = True
            univ_pg.add_run(f' The Test performed was a ')
            univ_pg.add_run(f'{univ_parameters["Test"]}').bold = True
            univ_pg.add_run(f' considering the variance between classes ')

            # Name of the string file to save tables and volcano plot as
            test_performed = univ_parameters["Test"].split(' ')[0]
            filename_string_abv = f'{univ_parameters["Test Class"]}_vs_{univ_parameters["Control Class"]}_{test_performed}'
            filename_string_abv = filename_string_abv + f'_pvalue{univ_parameters["p-value"]}_FC'
            filename_string_abv = filename_string_abv + f'{univ_parameters["Fold Change Threshold"]}'

            # Continuing description and Filename
            if univ_parameters["Expected Equal Var."]:
                univ_pg.add_run(f'equal.').bold = True
                filename_string_abv = filename_string_abv + f'_equalvariance'
            else:
                univ_pg.add_run(f'not equal.').bold = True
                filename_string_abv = filename_string_abv + f'_notequalvariance'
            univ_pg.add_run(f' Features were considered significant if the univariate test p-value was below the chosen ')
            univ_pg.add_run(f'{univ_parameters["p-value"]}').bold = True
            univ_pg.add_run(f' and the fold change (FC) between the average normalized intensity values (after missing value')
            univ_pg.add_run(f' imputation) of the 2 tested classes was greater than the chosen ')
            univ_pg.add_run(f'{univ_parameters["Fold Change Threshold"]}').bold = True
            univ_pg.add_run(f' in favour of either class. Pre-Treatment previously selected was repeated considering')
            univ_pg.add_run(f' only the samples belonging to the 2 classes selected before univariate analysis.')

            # Note about Fold Change threshold
            univ_pg2 = document.add_paragraph('Fold Change (FC) calculation is greatly affected by missing values. ')
            univ_pg2.add_run('Thus, considering the very high missing value occurrence in FT-ICR-MS data, FC results')
            univ_pg2.add_run(' should be taken with a grain of salt. The more missing values, ')
            univ_pg2.add_run('the less reliable FC and univariate test results are.')

            # Saving Univariate Test Results
            filt_filename_string = '/Report_Univar_res_' + filename_string_abv + '.xlsx'
            non_filt_filename_string = '/Report_Univar_nonfilt_res_' + filename_string_abv + '.xlsx'
            UnivarA_Store.univariate_results.to_excel(folder + filt_filename_string)
            pd.concat((UnivarA_Store.univariate_results_non_filt,
                DataFrame_Store.metadata_df.loc[UnivarA_Store.univariate_results_non_filt.index]),
                axis=1).to_excel(folder + non_filt_filename_string)

            # Univariate Test results
            univ_pg3 = document.add_paragraph('Univariate Analysis results: ')
            univ_pg3.add_run(f'{UnivarA_Store.univariate_results.shape[0]}').bold = True
            univ_pg3.add_run(f' metabolites were significant, ')
            univ_pg3.add_run(f'{sum(UnivarA_Store.univariate_results["Has Match?"])}').bold = True
            univ_pg3.add_run(f' of which are annotated. Volcano Plot of data shown below and univariate test results ')
            univ_pg3.add_run(f"and metabolic feature metadata with only the significant metabolites and with all metabolites")
            univ_pg3.add_run(f'are respectively stored in ')
            univ_pg3.add_run(f"'{filt_filename_string[1:]}' and '{non_filt_filename_string[1:]}'.")

            # Saving Volcano Plot figure and adding it to the document
            UnivarA_Store.Volcano_fig[0].write_image(folder+'/Report_VolcanoPlot - '+filename_string_abv+'.png', scale=4)
            UnivarA_Store.Volcano_fig[0].write_html(folder+'/Report_VolcanoPlot - '+filename_string_abv+".html")
            # Adding figure
            document.add_picture(folder+'/Report_VolcanoPlot - '+filename_string_abv+'.png', width=Cm(12))

        # End of section
        document.add_page_break()



    # Data Diversity Visualization Section
    if len(np.intersect1d(['Van Krevelen', 'Kendrick Mass Defect', 'Chem. Comp. Series'], stat_methods)) != 0:
        document.add_heading('Data Diversity Visualization Analysis', level=2)

        # Van Krevelen Plot Section
        if 'Van Krevelen' in stat_methods:
            # Heading
            document.add_heading('Van Krevelen Plots', level=3)

            # In case analysis was not performed
            if type(dataviz_store.VanKrevelen_plot[0]) == str:
                document.add_paragraph(('Van Krevelen Plots were not computed or no column with Formula annotations was'
                                        ' provided. Thus, this section will be skipped.'))

            else:
                # Description of the Van Krevelen Plots
                vk_pg = document.add_paragraph('Van Krevelen Plots were computed considering the formulas annotated in the following columns: ')
                vk_pg.add_run(', '.join(dataviz_store.vk_formula_to_consider)).bold = True
                vk_pg.add_run(('. If multiple formulas were assigned to a metabolic feature, they are both considered and '
                            'plotted. Multiple features with the same formula lead to multiple points in the same part of'
                            ' the Van Krevelen Plots. Below a Van Krevelen Plot is shown for each biological class in the '
                            'target considering only the metabolic features that appear in the samples of that class.'))

                # See what the points are highlighted by if anything
                if dataviz_store.vk_highlight_by != 'None':
                    size_colour_none = False
                    if dataviz_store.vk_colour:
                        if dataviz_store.vk_size:
                            vk_pg.add_run(f' Points colour and size are based on the ')
                        else:
                            vk_pg.add_run(f' Points colour are based on the ')
                    else:
                        if dataviz_store.vk_size:
                            vk_pg.add_run(f' Points size are based on the ')
                        else:
                            size_colour_none = True

                    # If indeed they are highlighted in size or colours, explain how
                    if not size_colour_none:
                        vk_pg.add_run(f'{dataviz_store.vk_highlight_by} of their average intensity value compared to others.')
                        if dataviz_store.vk_highlight_by != 'Rank':
                            vk_pg.add_run(' logInt represents the logarithm (base 10) of the average intensities.')

                        # Midpoint Sentence
                        vk_pg.add_run((f' The midpoint was set to {dataviz_store.vk_midpoint}, that is, '
                                    f'{dataviz_store.vk_midpoint*100} % of points were on the "low" intensity section of '
                                    'the scale and the remaining on the "high" intensity section of the scale.'))

                # Saving and adding each Van Krevelen Plot to the document
                for i in range(len(dataviz_store.VanKrevelen_plot)):
                    # Saving the figures
                    dataviz_store.VanKrevelen_plot[i].write_image(
                        folder+'/Report_'+dataviz_store.VanKrevelen_filenames[i]+'.png', scale=4)
                    dataviz_store.VanKrevelen_plot[i].write_html(
                        folder+'/Report_'+dataviz_store.VanKrevelen_filenames[i]+".html")
                    # Adding figure
                    document.add_picture(folder+'/Report_'+dataviz_store.VanKrevelen_filenames[i]+'.png', width=Cm(15))

                # End of section
                document.add_page_break()


        # Kendrick Mass Defect Plot Section
        if 'Kendrick Mass Defect' in stat_methods:
            # Heading
            document.add_heading('Kendrick Mass Defect Plots', level=3)

            # In case analysis was not performed
            if type(dataviz_store.KendrickMD_plot[0]) == str:
                document.add_paragraph(('Kendrick Mass Defect Plots were not computed or no column with feature Mass was'
                                        ' provided/found. Thus, this section will be skipped.'))

            else:
                # Description of the Kendrick Mass Defect Plots
                kmd_pg = document.add_paragraph('Kendrick Mass Defect Plots were computed considering the Mass column: ')
                kmd_pg.add_run(radiobox_neutral_mass.value).bold = True
                kmd_pg.add_run('. Masses were rounded ')

                # See how the masses were rounded
                if dataviz_store.kmd_mass_rounding == 'Up':
                    kmd_pg.add_run('Up').bold = True
                    kmd_pg.add_run(' to the next integer (thus, all Mass Defects will vary between 0 and 1).')
                else:
                    kmd_pg.add_run('to the ')
                    kmd_pg.add_run('Nearest').bold = True
                    kmd_pg.add_run(' integer (thus, all Mass Defects will vary between -0.5 and 0.5).')

                # If points were coloured by chemical composition series
                if len(dataviz_store.kmd_formula_to_consider) != 0:
                    kmd_pg.add_run(' Chemical formulas in the columns ')
                    kmd_pg.add_run(', '.join(dataviz_store.kmd_formula_to_consider)).bold = True
                    kmd_pg.add_run((' were considered to colours the points based on the chemical composition series they'
                                ' belong to (if found). When, for one metabolic feature, there are multiple candidate '
                                ' formulas and they do not belong to the same chemical composition series, they get '
                                    'assigned as Ambiguous.'))

                # Saving and adding each Kendrick Mass Defect Plot to the document
                for i in range(len(dataviz_store.KendrickMD_plot)):
                    # Saving the figures
                    dataviz_store.KendrickMD_plot[i].write_image(
                        folder+'/Report_'+dataviz_store.KendrickMD_filenames[i]+'.png', scale=4)
                    dataviz_store.VanKrevelen_plot[i].write_html(
                        folder+'/Report_'+dataviz_store.KendrickMD_filenames[i]+".html")
                    # Adding figure
                    document.add_picture(folder+'/Report_'+dataviz_store.KendrickMD_filenames[i]+'.png', width=Cm(15))

                # End of section
                document.add_page_break()


        # Chemical Composition Series Plot Section
        if 'Chem. Comp. Series' in stat_methods:
            # Heading
            document.add_heading('Chemical Composition Series Plot', level=3)

            # In case analysis was not performed
            if type(dataviz_store.CCS_plot[0]) == str:
                document.add_paragraph(('Chemical Composition Series Plot was not computed or no column with Formula annotations was'
                                        ' provided. Thus, this section will be skipped.'))

            else:
                # Description of the Chemical Composition Series Plot
                ccs_pg = document.add_paragraph(("Chemical Composition Series Plot was "
                                                'computed considering the formulas in the following columns: '))
                ccs_pg.add_run(', '.join(dataviz_store.ccs_formula_to_consider)).bold = True
                ccs_pg.add_run(('. If multiple formulas were assigned to a metabolic feature, each one was counted in this'
                            ' plot. That is, if a peak in a class has 3 possible candidate formulas, 2 belonging to the '
                                "'CHO' series and another to the 'CHOP' series; then 2 formulas will be added to the 'CHO' "
                                "series and 1 to the 'CHOP' series. Thus, we are considering that the 3 elementary formulas "
                                " are represented by that feature (probably an overestimation). To provide an idea of how "
                                "extensive this effect is, a description is also provided detailing how many formulas are "
                                "being considered for each class and from how many different features (m/z peaks) they came from."))

                # Description of considered formulas and from how many metabolic features they come from
                document.add_paragraph(dataviz_store.ccs_desc.replace('<br />', '\n').replace('**', ''))

                # Creating filename for the Chemical Composition Series Plot and saving it
                filename_string = '/Report_CCS_Plot_formulacolumns'
                for cl in dataviz_store.ccs_formula_to_consider:
                    filename_string = filename_string + f'_{cl}'
                dataviz_store.CCS_plot[0].write_image(folder+filename_string+'.png', scale=4)
                dataviz_store.CCS_plot[0].write_html(folder+filename_string+".html")

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(12.5))

                # Table with the counts of each series for each class
                table_ccs_counts = document.add_table(rows=len(dataviz_store.ccs_df.index)+1,
                                                            cols=len(dataviz_store.ccs_df.columns)+1,
                                                            style='Light Grid')
                table_ccs_counts = fill_word_table(table_ccs_counts, dataviz_store.ccs_df)

                # Give space after table
                document.add_paragraph()

            # End of section
            document.add_page_break()



    # HMDB Pathway Assignment Section
    if 'HMDB Pathway Assignment' in stat_methods:
        # Heading
        document.add_heading('HMDB Pathway Assignment', level=2)

        # In case analysis was not performed
        if len(PathAssign_store.current_hmdb_id_cols) == 0:
            document.add_paragraph(('HMDB Pathway Assignment to HMDB IDs was not performed or no column with HMDB IDs was'
                                        ' selected. Thus, this section will be skipped.'))

        else:
            # Descriptions of HMDB Pathway Assignment Procedure
            path_pg = document.add_paragraph('This section searches the HMDB IDs provided against a database containing ')
            path_pg.add_run(('55872 HMDB compound IDs that have at least 1 pathway associated. The file was created based '
                            'on the RAMP database (https://rampdb.nih.gov/) that attempts to aggregate the HMDB, Reactome, '
                            'WikiPathways and KEGG databases (https://academic.oup.com/bioinformatics/article/39/1/btac726/6827287'
                            "). HMDB IDs, e.g. 'HMDB0000001', should always start with 'HMDB' followed by 7 numbers. IDs "
                            "provided that do not follow this formatting will be disregarded and not considered as "
                            "'HMDB-like IDs'."))

            # Descriptions of HMDB Pathway Assignment Performed
            path_pg2 = document.add_paragraph('The columns selected to search for HMDB-like IDs were the following: ')
            path_pg2.add_run(', '.join(PathAssign_store.current_hmdb_id_cols)).bold = True
            path_desc_split = PathAssign_store.assign_desc.split('**')
            path_pg2.add_run('. From this/these columns, ')
            path_pg2.add_run(path_desc_split[1]).bold = True
            path_pg2.add_run(' possible HMDB IDs were found, ')
            path_pg2.add_run(path_desc_split[3]).bold = True
            path_pg2.add_run(' of which were HMDB-like IDs. From those, ')
            path_pg2.add_run(path_desc_split[5]).bold = True
            path_pg2.add_run(path_desc_split[6])

            # Saving the pathway assignments DataFrame
            # Seeing indexes with HMDB-like IDs
            considered_idxs = []
            for idx in PathAssign_store.pathway_assignments.index:
                if len(idx) == 11:
                    if idx.startswith('HMDB'):
                        try:
                            int(idx[4:])
                            considered_idxs.append(idx)
                        except:
                            continue
            # Building the datafile name
            filename_string = f'/Report_HMDB_IDs_PathwaysMatching_to_onlyHMDB-likeIDs_in'
            for col in PathAssign_store.current_hmdb_id_cols:
                filename_string = filename_string + f'_{col}'
            filename_string = filename_string + '.xlsx'
            # Saving the file
            PathAssign_store.pathway_assignments.loc[considered_idxs].to_excel(folder+filename_string)

            path_pg2.add_run(f"'{filename_string[1:]}' excel contains the results of the pathway matching made ")
            path_pg2.add_run("only").bold = True
            path_pg2.add_run(f" showing the IDs considered HMDB-like found in the provided columns.")

        # End of section
        document.add_page_break()



    # BinSim Analysis Section
    if len(np.intersect1d(['BinSim - PCA', 'BinSim - HCA', 'BinSim - PLS-DA', 'BinSim - RF'], stat_methods)) != 0:
        document.add_heading('Binary Simplification (BinSim) Treated Dataset - Multivariate Analysis', level=2)

        binsim_pg = document.add_paragraph(f'Binary Simplification (BinSim) treated data consists of transforming the ')
        binsim_pg.add_run('original data matrix by changing all intensity values detected to 1 and all missing values to 0. ')
        binsim_pg.add_run('Thus, it represents the metabolic feature occurrence information complementary to the standard ')
        binsim_pg.add_run('intensity-based analysis.')

        # PCA Section - BinSim version
        if 'BinSim - PCA' in stat_methods:
            # Heading
            document.add_heading('Principal Component Analysis (PCA) - BinSim Version', level=3)

            # If there is a PCA Projection
            if type(PCA_params_binsim.PCA_plot[0]) != str:
                # Description
                binsim_pca_pg = document.add_paragraph(f'Principal Component Analysis on BinSim treated data was made with ')
                binsim_pca_pg.add_run(f'{len(PCA_params_binsim.controls.widgets["PCx"].options)} components.')

                binsim_pca_pg2 = document.add_paragraph()
                if PCA_params_binsim.n_dimensions == '2 Components':
                    binsim_pca_pg2.add_run(f'2-D PCA Projection Plot of Components: {PCA_params_binsim.PCx} ').bold = True
                    binsim_pca_pg2.add_run(f'and {PCA_params_binsim.PCy} - BinSim').bold = True
                else:
                    binsim_pca_pg2.add_run(f'3-D PCA Projection Plot of Components: {PCA_params_binsim.PCx}, ').bold = True
                    binsim_pca_pg2.add_run(f'{PCA_params_binsim.PCy} and {PCA_params_binsim.PCz} - BinSim').bold = True
                binsim_pca_pg2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Creating filename for the PCA Projection and saving it
                filename_string = '/Report_BinSim_PCA_plot'
                if PCA_params_binsim.ellipse_draw:
                    if PCA_params_binsim.confidence != 0:
                        filename_string = filename_string + f'_ellipse({PCA_params_binsim.confidence*100}%confidence)'
                    else:
                        filename_string = filename_string + f'_ellipse({PCA_params_binsim.confidence_std}std)'
                PCA_params_binsim.PCA_plot[0].write_image(folder+filename_string+'.png', scale=4)
                PCA_params_binsim.PCA_plot[0].write_html(folder+filename_string+".html")

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(12.5))

                # If there is an explained variance plot
                if type(PCA_params_binsim.PCA_plot[0]) != str:
                    # Description
                    binsim_pca_pg3 = document.add_paragraph()
                    binsim_pca_pg3.add_run('Cumulative Explained Variance (by Principal Component) Plot - BinSim').bold = True
                    binsim_pca_pg3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Creating filename for the PCA Projection and saving it
                    filename_string = '/Report_BinSim_PCA_exp_var_plot'
                    PCA_params_binsim.exp_var_fig_plot[0].write_image(folder+filename_string+'.png', scale=4)
                    PCA_params_binsim.exp_var_fig_plot[0].write_html(folder+filename_string+".html")

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(12.5))

                # If there is an explained variance plot
                if type(PCA_params_binsim.scatter_PCA_plot[0]) != str:
                    # Description
                    binsim_pca_pg4 = document.add_paragraph()
                    binsim_pca_pg4.add_run('2-D PCA Projection (Scatter) Plot of the Principal Components - BinSim').bold = True
                    binsim_pca_pg4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Creating filename for the PCA Projection and saving it
                    filename_string = '/Report_BinSim_PCA_scatter_plot'
                    PCA_params_binsim.scatter_PCA_plot[0].write_image(folder+filename_string+'.png', scale=4)
                    PCA_params_binsim.scatter_PCA_plot[0].write_html(folder+filename_string+".html")

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(15))

            # If there is no PCA analysis performed on BinSim treated data
            else:
                binsim_pca_pg = document.add_paragraph('Principal Component Analysis was not performed with the BinSim ')
                binsim_pca_pg.add_run(f"treated data during the Data Analysis. Thus, this section will be skipped.")


        # HCA Section - BinSim version
        if 'BinSim - HCA' in stat_methods:
            # Heading
            document.add_heading('Hierarchical Clustering Analysis (HCA) - BinSim Version', level=3)

            # If there is a dendrogram
            if type(HCA_params_binsim.HCA_plot[0]) != str:
                # Description
                binsim_hca_pg = document.add_paragraph(f'Hierarchical Clustering Analysis on BinSim treated data was made ')
                binsim_hca_pg.add_run(f'with {HCA_params_binsim.dist_metric} distance metric and ')
                binsim_hca_pg.add_run(f'{HCA_params_binsim.link_metric} linkage metric.')

                # Creating filename for the dendrogram and saving it
                filename = f'/Report_BinSim_HCA_plot_{HCA_params_binsim.dist_metric}Dist_{HCA_params_binsim.link_metric}Linkage'
                HCA_params_binsim.HCA_plot[0].savefig(folder + filename_string + '.png', dpi=HCA_params_binsim.dpi)

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(12))

        # End of unsupervised analysis of BinSim section
        document.add_page_break()


        # PLS-DA Section - BinSim Version
        if 'BinSim - PLS-DA' in stat_methods:
            # Heading
            document.add_heading('Partial Least Squares - Discriminant Analysis (PLS-DA) - BinSim Version', level=3)

            # If there is a PLS optimization
            if type(PLSDA_store_binsim.optim_figure[0]) != str:
                # Minor Heading
                document.add_heading('Number of Components Optimization - BinSim', level=4)

                # Description
                n_min_components, n_max_components = PLSDA_store_binsim.current_other_plsda_params["n_min_max_components"]
                binsim_plsda_pg = document.add_paragraph(f'First, an optimization of the number of components to use to fit ')
                binsim_plsda_pg.add_run(f'a PLS-DA model on BinSim treated data was made by fitting PLS models from ')
                binsim_plsda_pg.add_run(f'{n_min_components} to {n_max_components} components and evaluating their Q2 ')
                binsim_plsda_pg.add_run(f'(mainly) and R2 scores estimated by {PLSDA_store_binsim.current_other_plsda_params["n_fold_optim"]}')
                binsim_plsda_pg.add_run(f'-fold stratified cross-validation. Results are shown in the figure below with a ')
                binsim_plsda_pg.add_run(f'maximum Q2 value with {PLSDA_store_binsim.rec_components} components.')

                # Plot title
                binsim_plsda_pg2 = document.add_paragraph()
                binsim_plsda_pg2.add_run('PLS-DA Component Optimization Plot - BinSim').bold = True
                binsim_plsda_pg2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Creating filename for the PLS optimization plot and saving it
                filename_string = f'/Report_BinSim_PLS_optim_plot_{n_min_components}to{n_max_components}'
                filename_string = filename_string + f'components_{PLSDA_store_binsim.current_other_plsda_params["n_fold_optim"]}-'
                filename_string = filename_string + f'foldstratCV_scale{PLSDA_store_binsim.current_other_plsda_params["scale_optim"]}'
                PLSDA_store_binsim.optim_figure[0].write_image(folder+filename_string+'.png', scale=4)
                PLSDA_store_binsim.optim_figure[0].write_html(folder+filename_string+".html")

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(15))

            else:
                document.add_paragraph(f'Optimization of the number of components to fit the PLS-DA model on BinSim treated data was not performed.')


            # If the PLS-DA model was fitted
            if type(PLSDA_store_binsim.models[0]) != str:
                # Minor Heading
                document.add_heading('PLS-DA Model Fitting and Results - BinSim', level=4)

                # Description
                binsim_plsda_params = PLSDA_store_binsim.current_plsda_params
                binsim_plsda_pg3 = document.add_paragraph(f'PLS-DA model was fitted on BinSim treated data using ')
                binsim_plsda_pg3.add_run(f'{binsim_plsda_params["n_components"]} components estimating model performance and')
                binsim_plsda_pg3.add_run(f' feature importance ({binsim_plsda_params["feat_imp"]} score method) by ')
                binsim_plsda_pg3.add_run(f'{binsim_plsda_params["n_folds"]}-fold stratified cross-validation repeated ')
                binsim_plsda_pg3.add_run(f'{binsim_plsda_params["n_iterations"]} times (randomized folds in cross-validation)')
                binsim_plsda_pg3.add_run(f'. Model Performance results are shown in the table below (by the metrics chosen) ')

                # Building the filename and saving PLS-DA Feature Importance Table
                filename_string = f'/Report_BinSim_PLS-DA_FeatImp_{binsim_plsda_params["feat_imp"]}_model_params_components'
                filename_string = filename_string + f'{binsim_plsda_params["n_components"]}_{binsim_plsda_params["n_folds"]}'
                filename_string = filename_string + f'-foldstratCV_iterations{binsim_plsda_params["n_iterations"]}_scale'
                filename_string = filename_string + f'{binsim_plsda_params["scale"]}.xlsx'
                PLSDA_store_binsim.feat_impor.to_excel(folder+filename_string)

                # Finishing up the description
                binsim_plsda_pg3.add_run(f"and Feature Importance Table was saved as '{filename_string[1:]}'. ")
                binsim_plsda_pg3.add_run(f'Furthermore, below we also show a PLS Projection of chosen Latent Variables.')

                # Table with results of PLS-DA
                table_plsda_results_binsim = document.add_table(rows=len(PLSDA_store_binsim.n_results.index)+1,
                                                            cols=len(PLSDA_store_binsim.n_results.columns)+1,
                                                            style='Light Grid')
                table_plsda_results_binsim = fill_word_table(table_plsda_results_binsim, PLSDA_store_binsim.n_results)

                # Give space after table
                document.add_paragraph()

                binsim_plsda_pg4 = document.add_paragraph()
                if PLSDA_store_binsim.n_dimensions == '2 Components':
                    binsim_plsda_pg4.add_run(f'2-D PLS Projection Plot of Latent Variables: {PLSDA_store_binsim.LVx}').bold = True
                    binsim_plsda_pg4.add_run(f' and {PLSDA_store_binsim.LVy} - BinSim').bold = True
                else:
                    binsim_plsda_pg4.add_run(f'3-D PCA Projection Plot of Latent Variables: {PLSDA_store_binsim.LVx}').bold = True
                    binsim_plsda_pg4.add_run(f', {PLSDA_store_binsim.LVy} and {PLSDA_store_binsim.LVz} - BinSim').bold = True
                binsim_plsda_pg4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Creating filename for the PLS Projection and saving it
                filename_string = '/Report_BinSim_PLS_projection_plot'
                if PLSDA_store_binsim.n_dimensions == '2 Components':
                    if PLSDA_store_binsim.ellipse_draw:
                        if PLSDA_store_binsim.confidence != 0:
                            filename_string = filename_string + f'_ellipse({PLSDA_store_binsim.confidence*100}%confidence)'
                        else:
                            filename_string = filename_string + f'_ellipse({PLSDA_store_binsim.confidence_std}std)'
                PLSDA_store_binsim.PLS_plot[0].write_image(folder+filename_string+'.png', scale=4)
                PLSDA_store_binsim.PLS_plot[0].write_html(folder+filename_string+".html")

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(15))


                # If there is Permutation Test plot
                if type(PLSDA_store_binsim.perm_figure[0]) != str:
                    # Minor Heading
                    document.add_heading('PLS-DA Permutation Test - BinSim', level=4)

                    # Description
                    binsim_plsda_pg5 = document.add_paragraph('Permutation Test was performed on BinSim treated data with ')
                    binsim_plsda_pg5.add_run(f"the same number of components as before - {PLSDA_store_binsim.current_plsda_params_permutation['n_components']}")
                    binsim_plsda_pg5.add_run(f" - and with {PLSDA_store_binsim.current_plsda_params_permutation['n_permutations']}")
                    binsim_plsda_pg5.add_run(f" permutations. Model performance was evaluated by ")
                    binsim_plsda_pg5.add_run(f"{PLSDA_store_binsim.current_plsda_params_permutation['perm_metric']} ")
                    binsim_plsda_pg5.add_run(f"estimated with {PLSDA_store_binsim.current_plsda_params_permutation['n_folds']}")
                    binsim_plsda_pg5.add_run(f"-fold stratified cross-validation.")

                    # Creating filename for the PLS-DA Permutation Test and saving it
                    filename_string = f'/Report_BinSim_PLS-DA_permutation_test_{PLSDA_store_binsim.current_plsda_params_permutation["n_permutations"]}perm_'
                    filename_string = filename_string + f'{PLSDA_store_binsim.current_plsda_params_permutation["n_components"]}comp_'
                    filename_string = filename_string + f'{PLSDA_store_binsim.current_plsda_params_permutation["n_folds"]}-foldstratCV_scale'
                    filename_string = filename_string + f'{PLSDA_store_binsim.current_plsda_params_permutation["scale"]}_metric'
                    filename_string = filename_string + f'{PLSDA_store_binsim.current_plsda_params_permutation["perm_metric"]}'
                    PLSDA_store_binsim.perm_figure[0].savefig(folder+filename_string+'.png', dpi=PLSDA_store_binsim.dpi)

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(15))


                # If there is ROC Curve plot
                if type(PLSDA_store_binsim.ROC_figure[0]) != str:
                    # Minor Heading
                    document.add_heading('PLS-DA ROC Curve - BinSim', level=4)

                    # Description
                    binsim_roc_params = PLSDA_store_binsim.current_other_plsda_params['ROC_filename'].split('_')
                    if len(target_list.classes) > 2:
                        binsim_plsda_pg6 = document.add_paragraph('Since there are more than 2 classes, ROC curves were computed')
                        binsim_plsda_pg6.add_run(f" with a 1vsAll scheme for each of the {len(target_list.classes)} classes.")
                    else:
                        binsim_plsda_pg6 = document.add_paragraph('Since there are only 2 classes, ROC curves were computed considering ')
                        binsim_plsda_pg6.add_run(f"{binsim_roc_params[2][:-8]}").bold = True
                        binsim_plsda_pg6.add_run(f" as the positive class.")
                    binsim_plsda_pg6.add_run(f' Other parameters were maintained: number of components - {binsim_roc_params[3][:-10]}')
                    binsim_plsda_pg6.add_run(f"; {binsim_roc_params[4][:-7]} stratified cross-validation; scale - ")
                    binsim_plsda_pg6.add_run(f"{binsim_roc_params[6][5:]}. Finally, ROC Curve estimation was repeated ")
                    binsim_plsda_pg6.add_run(f"{binsim_roc_params[5][:-10]} times (randomized cross-validation folds).")

                    # Creating filename for the PLS-DA ROC Curve and saving it
                    filename_string = '/Report_BinSim_' + PLSDA_store_binsim.current_other_plsda_params['ROC_filename'][:-7]
                    PLSDA_store_binsim.ROC_figure[0].write_image(folder+filename_string+'.png', scale=4)
                    PLSDA_store_binsim.ROC_figure[0].write_html(folder+filename_string+".html")

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(15))

            # If there is no PLS-DA model fitted on the BinSim treated data
            else:
                binsim_plsda_pg3 = document.add_paragraph('PLS-DA model was not fitted on the BinSim treated ')
                binsim_plsda_pg3.add_run(f"data during the Data Analysis. Thus, this section will be skipped.")


        # Random Forest Section - BinSim Version
        if 'BinSim - Random Forest' in stat_methods:
            # Heading
            document.add_heading('Random Forest (RF) - BinSim', level=3)

            # If there is a Random Forest Tree Number optimization
            if type(RF_store_binsim.optim_figure[0]) != str:
                # Minor Heading
                document.add_heading('Number of Trees Optimization - BinSim', level=4)

                # Description
                n_min_trees, n_max_trees = RF_store_binsim.current_other_rf_params["n_min_max_trees"]
                binsim_rf_pg = document.add_paragraph(f'First, an optimization of the number of trees to use to fit a ')
                binsim_rf_pg.add_run(f'Random Forest model on BinSim treated data was made by fitting RF models from ')
                binsim_rf_pg.add_run(f'{n_min_trees} to {n_max_trees} trees in {RF_store_binsim.current_other_rf_params["n_interval"]}')
                binsim_rf_pg.add_run(f' tree steps, evaluating their model accuracy estimated by ')
                binsim_rf_pg.add_run(f'{RF_store_binsim.current_other_rf_params["n_fold_optim"]}-fold stratified')
                binsim_rf_pg.add_run(f' cross-validation. Results are shown in the figure below.')

                # Plot title
                binsim_rf_pg2 = document.add_paragraph()
                binsim_rf_pg2.add_run('Random Forest Number of Trees Optimization Plot - BinSim').bold = True
                binsim_rf_pg2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Creating filename for the Random Forest optimization plot and saving it
                filename_string = f'/Report_BinSim_RF_optim_plot_{RF_store_binsim.current_other_rf_params["n_fold_optim"]}'
                filename_string = filename_string + f'-foldStratCV_{n_min_trees}to{n_max_trees}trees_'
                filename_string = filename_string + f'({RF_store_binsim.current_other_rf_params["n_interval"]}interval)'
                RF_store_binsim.optim_figure[0].write_image(folder+filename_string+'.png', scale=4)
                RF_store_binsim.optim_figure[0].write_html(folder+filename_string+".html")

                # Adding figure
                document.add_picture(folder+filename_string+'.png', width=Cm(15))

            else:
                document.add_paragraph(f'Optimization of the number of trees to fit the Random Forest model on BinSim treated data was not performed.')


            # If the Random Forest model was fitted
            if type(RF_store_binsim.models[0]) != str:
                # Minor Heading
                document.add_heading('Random Forest Model Fitting and Results - BinSim', level=4)

                # Description
                binsim_rf_params = RF_store_binsim.current_rf_params
                binsim_rf_pg3 = document.add_paragraph(f'Random Forest model was fitted on BinSim treated data with ')
                binsim_rf_pg3.add_run(f'{binsim_rf_params["n_trees"]} trees estimating model performance and feature ')
                binsim_rf_pg3.add_run(f'importance by {binsim_rf_params["n_folds"]}-fold stratified cross-validation ')
                binsim_rf_pg3.add_run(f'repeated {binsim_rf_params["n_iterations"]} times (randomized folds in ')
                binsim_rf_pg3.add_run(f'cross-validation). Model Performance results are shown in the table below (by the ')

                # Building the filename and saving Random Forest Feature Importance Table
                filename_string = f'/Report_BinSim_RF_FeatImp_Gini_model_params_{binsim_rf_params["n_trees"]}trees_'
                filename_string = filename_string + f'{binsim_rf_params["n_folds"]}-foldstratCV_iterations'
                filename_string = filename_string + f'{binsim_rf_params["n_iterations"]}.xlsx'
                RF_store_binsim.feat_impor.to_excel(folder+filename_string)

                # Finishing up the description
                binsim_rf_pg3.add_run(f"metrics chosen) and Feature Importance Table was saved as '{filename_string[1:]}'.")

                # Table with results of Random Forest
                table_rf_results_binsim = document.add_table(rows=len(RF_store_binsim.n_results.index)+1,
                                                            cols=len(RF_store_binsim.n_results.columns)+1,
                                                            style='Light Grid')
                table_rf_results_binsim = fill_word_table(table_rf_results_binsim, RF_store_binsim.n_results)

                # Give space after table
                document.add_paragraph()


                # If there is Permutation Test plot
                if type(RF_store_binsim.perm_figure[0]) != str:
                    # Minor Heading
                    document.add_heading('Random Forest Permutation Test - BinSim', level=4)

                    # Description
                    binsim_rf_params_perm = RF_store_binsim.current_rf_params_permutation
                    binsim_rf_pg4 = document.add_paragraph('Permutation Test was performed on the BinSim treated data with ')
                    binsim_rf_pg4.add_run(f"the same number of trees as before - {binsim_rf_params_perm['n_trees']} - and ")
                    binsim_rf_pg4.add_run(f"with {binsim_rf_params_perm['n_permutations']} permutations. Model performance ")
                    binsim_rf_pg4.add_run(f"was evaluated by {binsim_rf_params_perm['perm_metric']} estimated with ")
                    binsim_rf_pg4.add_run(f"{binsim_rf_params_perm['n_folds']}-fold stratified cross-validation.")

                    # Creating filename for the Random Forest Permutation Test and saving it
                    filename_string = f'/Report_BinSim_RF_permutation_test_{binsim_rf_params_perm["n_permutations"]}perm_'
                    filename_string = filename_string + f'{binsim_rf_params_perm["n_trees"]}comp_'
                    filename_string = filename_string + f'{binsim_rf_params_perm["n_folds"]}-foldstratCV_metric'
                    filename_string = filename_string + f'{binsim_rf_params_perm["perm_metric"]}'
                    RF_store_binsim.perm_figure[0].savefig(folder+filename_string+'.png', dpi=RF_store_binsim.dpi)

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(15))


                # If there is ROC Curve plot
                if type(RF_store_binsim.ROC_figure[0]) != str:
                    # Minor Heading
                    document.add_heading('Random Forest ROC Curve - BinSim', level=4)

                    # Description
                    binsim_roc_params = RF_store_binsim.current_other_rf_params['ROC_filename'].split('_')
                    if len(target_list.classes) > 2:
                        binsim_rf_pg5 = document.add_paragraph('Since there are more than 2 classes, ROC curves were computed')
                        binsim_rf_pg5.add_run(f" with a 1vsAll scheme for each of the {len(target_list.classes)} classes.")
                    else:
                        binsim_rf_pg5 = document.add_paragraph('Since there are only 2 classes, ROC curves were computed considering ')
                        binsim_rf_pg5.add_run(f"{binsim_roc_params[2][:-8]}").bold = True
                        binsim_rf_pg5.add_run(f" as the positive class.")
                    binsim_rf_pg5.add_run(f' Other parameters were maintained: number of trees - {binsim_roc_params[3][:-5]}; ')
                    binsim_rf_pg5.add_run(f"{binsim_roc_params[4][:-7]} stratified cross-validation. Finally, ROC Curve estimation was ")
                    binsim_rf_pg5.add_run(f"repeated {binsim_roc_params[5][:-10]} times (randomized cross-validation folds).")

                    # Creating filename for the Random Forest ROC Curve and saving it
                    filename_string = '/Report_BinSim_' + RF_store_binsim.current_other_rf_params['ROC_filename'][:-7]
                    RF_store_binsim.ROC_figure[0].write_image(folder+filename_string+'.png', scale=4)
                    RF_store_binsim.ROC_figure[0].write_html(folder+filename_string+".html")

                    # Adding figure
                    document.add_picture(folder+filename_string+'.png', width=Cm(15))

            # If there is no Random Forest model fitted
            else:
                binsim_rf_pg3 = document.add_paragraph('Random Forest model was not fitted on the BinSim treated ')
                binsim_rf_pg3.add_run(f"data during the Data Analysis. Thus, this section will be skipped.")

        # End of BinSim section
        document.add_page_break()




    # Saving the document if possible
    try:
        document.save(folder+'/Report.docx')
        os.chmod(folder+'/Report.docx', S_IREAD|S_IRGRP|S_IROTH)
    except:
        pn.state.notifications.error(f'Report.docx could not be saved since it already existed in {folder} folder.')
        while len(rep_gen_page) > 6:
            rep_gen_page.pop(5)
        raise ValueError(f'Report.docx could not be saved since it already existed in {folder} folder.')




# Saving Parameters function
def save_parameters(filename, RepGen, UnivarA_Store, n_databases, adducts_to_search_widget, DB_dict, checkbox_com_exc,
                    com_exc_compounds, PCA_params, HCA_params, PLSDA_store, RF_store, dataviz_store, PCA_params_binsim,
                    HCA_params_binsim, PLSDA_store_binsim, RF_store_binsim, include_data_analysis=True):
    "Creates a json file containing relevant parameters used in current dataset analysis."

    # Dict with parameters to be saved
    params_to_be_saved = {}

    # Saving Data Reading Related Parameters (just to check for the adducts later on)
    params_to_be_saved['Data Reading'] = {'type_of_mass_values_in_file': RepGen.type_of_mass_values_in_file}

    # Saving Data Filtering Related Parameters
    params_to_be_saved['Data Filtering'] = {'filt_method': UnivarA_Store.filt_method,
                                           'filt_kw': UnivarA_Store.filt_kw} # Special Precautions needed when loading


    # Saving Data Annotation Related Parameters
    # Main Parameters
    params_to_be_saved['Data Annotation'] = {'n_databases': n_databases.value,
                                            'annotation_margin_method': RepGen.annotation_margin_method,
                                            'annotation_ppm_deviation': RepGen.annotation_margin_ppm_deviation,
                                            'annotation_Da_deviation': RepGen.annotation_margin_Da_deviation,
                                            'adducts_to_be_searched': adducts_to_search_widget.value}
    # Database Specific Parameters
    for d in range(n_databases.value):
        params_to_be_saved['Data Annotation'][d] = {'file': DB_dict[str(d+1)].file,
                                                   'abv': DB_dict[str(d+1)].abv,
                                                   'IDcol': DB_dict[str(d+1)].IDcol,
                                                   'annotation': DB_dict[str(d+1)].annotation,
                                                   'formula': DB_dict[str(d+1)].formula,}


    # Saving Data Pre-Treatment Related Parameters
    params_to_be_saved['Data Pre-Treatment'] = {'mvi_method': UnivarA_Store.mvi_method,
                                               'mvi_kw': UnivarA_Store.mvi_kw,
                                               'norm_method': UnivarA_Store.norm_method, # Special Precautions needed when loading
                                               'norm_kw': UnivarA_Store.norm_kw, # Special Precautions needed when loading
                                               'tf_method': UnivarA_Store.tf_method,
                                               'tf_kw': UnivarA_Store.tf_kw,
                                               'scaling_method': UnivarA_Store.scaling_method,
                                               'scaling_kw': UnivarA_Store.scaling_kw}
    if UnivarA_Store.tf_method == 'None':
        params_to_be_saved['Data Pre-Treatment']['tf_method'] = 'None'


    # Whether or not to include parameters related to data analysis:
    if include_data_analysis:
        # Saving Common and Exclusive Compound Analysis Related Paramaeters
        params_to_be_saved['Com. and Exc. Comp.'] = {# General
                                                    'chosen': checkbox_com_exc.value, # Technically can be different
                                                    # Venn Diagram Related
                                                    'venn_alpha': com_exc_compounds.venn_alpha,
                                                    'type_of_venn': com_exc_compounds.type_of_venn,
                                                    'dpi_venn': com_exc_compounds.dpi_venn,
                                                    # Intersection Plot Related
                                                    'inter_include_counts_percentages': com_exc_compounds.inter_include_counts_percentages,
                                                    'dpi_inter': com_exc_compounds.dpi_inter}


        # Saving Unsupervised Analysis Related Parameters
        # PCA
        params_to_be_saved['PCA'] = {'n_components': PCA_params.n_components, # Remember to load the widget too
                                    'ellipse_draw': PCA_params.ellipse_draw,
                                    'confidence': PCA_params.confidence,
                                    'confidence_std': PCA_params.confidence_std,
                                    'dot_size': PCA_params.dot_size}
        # HCA
        params_to_be_saved['HCA'] = {'dist_metric': HCA_params.dist_metric,
                                    'link_metric': HCA_params.link_metric,
                                    'fig_x': HCA_params.fig_x,
                                    'fig_y': HCA_params.fig_y,
                                    'dpi': HCA_params.dpi,}


        # Saving Supervised Analysis Related Parameters
        # PLS-DA
        params_to_be_saved['PLS-DA'] = {# Optimization Specific Related
                                        'n_min_max_components': PLSDA_store.n_min_max_components, # Technically can be different
                                        # PLS-DA Model Fit
                                        'n_fold': PLSDA_store.n_fold,
                                        'scale': PLSDA_store.scale,
                                        'n_components': PLSDA_store.n_components,
                                        'n_iterations': PLSDA_store.n_iterations,
                                        'metrics_to_use': PLSDA_store.metrics_to_use, # Technically can be different
                                        'imp_feature_metric': PLSDA_store.imp_feature_metric,
                                        # PLS Projection Related
                                        'ellipse_draw': PLSDA_store.ellipse_draw,
                                        'confidence': PLSDA_store.confidence,
                                        'confidence_std': PLSDA_store.confidence_std,
                                        'dot_size': PLSDA_store.dot_size,
                                        # PLS-DA Permutation Test Related
                                        'n_perm': PLSDA_store.n_perm,
                                        'perm_metric': PLSDA_store.perm_metric,
                                        'dpi': PLSDA_store.dpi, # Technically can be different
                                        # PLS-DA ROC Curve Related
                                        'roc_n_iter': PLSDA_store.roc_n_iter} # Technically can be different
        # Update to actually used parameters if possible
        if 'n_folds' in PLSDA_store.current_plsda_params.keys():
            for k,v in PLSDA_store.current_plsda_params.items():
                if k == 'n_folds':
                    params_to_be_saved['PLS-DA']['n_fold'] = v
                elif k == 'feat_imp':
                    params_to_be_saved['PLS-DA']['imp_feature_metric'] = v
                else:
                    params_to_be_saved['PLS-DA'][k] = v
        if 'n_min_max_components' in PLSDA_store.current_other_plsda_params.keys():
            params_to_be_saved['PLS-DA']['n_min_max_components'] = PLSDA_store.current_other_plsda_params[
                'n_min_max_components']
        if 'ROC_filename' in PLSDA_store.current_other_plsda_params.keys():
            params_to_be_saved['PLS-DA']['roc_n_iter'] = PLSDA_store.current_other_plsda_params[
                'ROC_filename'].split('_')[5][:-10]
        if 'perm_metric' in PLSDA_store.current_plsda_params_permutation.keys():
            params_to_be_saved['PLS-DA']['n_perm'] = PLSDA_store.current_plsda_params_permutation['n_permutations']
            params_to_be_saved['PLS-DA']['perm_metric'] = PLSDA_store.current_plsda_params_permutation['perm_metric']

        # Random Forest
        params_to_be_saved['Random Forest'] = {# Optimization Specific Related
                                            'n_min_max_trees': RF_store.n_min_max_trees,
                                            'n_interval': RF_store.n_interval,
                                            # Random Forest Model Fit
                                            'n_fold': RF_store.n_fold,
                                            'n_trees': RF_store.n_trees,
                                            'n_iterations': RF_store.n_iterations,
                                            'metrics_to_use': RF_store.metrics_to_use, # Technically can be different
                                            # RF Permutation Test Related
                                            'n_perm': RF_store.n_perm,
                                            'perm_metric': RF_store.perm_metric,
                                            'dpi': RF_store.dpi, # Technically can be different
                                            # RF ROC Curve related
                                            'roc_n_iter': RF_store.roc_n_iter,}
        # Update to actually used parameters if possible
        if 'n_folds' in RF_store.current_rf_params.keys():
            for k,v in RF_store.current_rf_params.items():
                if k == 'n_folds':
                    params_to_be_saved['Random Forest']['n_fold'] = v
                else:
                    params_to_be_saved['Random Forest'][k] = v
        if 'n_min_max_trees' in RF_store.current_other_rf_params.keys():
            params_to_be_saved['Random Forest']['n_min_max_trees'] = RF_store.current_other_rf_params['n_min_max_trees']
            params_to_be_saved['Random Forest']['n_interval'] = RF_store.current_other_rf_params['n_interval']
        if 'ROC_filename' in RF_store.current_other_rf_params.keys():
            params_to_be_saved['Random Forest']['roc_n_iter'] = RF_store.current_other_rf_params[
                'ROC_filename'].split('_')[5][:-10]
        if 'perm_metric' in RF_store.current_rf_params_permutation.keys():
            params_to_be_saved['Random Forest']['n_perm'] = RF_store.current_rf_params_permutation['n_permutations']
            params_to_be_saved['Random Forest']['perm_metric'] = RF_store.current_rf_params_permutation['perm_metric']


        # Saving Univariate Analysis Related Parameters
        params_to_be_saved['Univariate Analysis'] = {'univariate_test': UnivarA_Store.univariate_test,
                                                    'expected_equal_variance': UnivarA_Store.expected_equal_variance,
                                                    'p_value_threshold': UnivarA_Store.p_value_threshold,
                                                    'fold_change_threshold': UnivarA_Store.fold_change_threshold,
                                                    'color_non_sig': UnivarA_Store.color_non_sig,
                                                    'color_down_sig': UnivarA_Store.color_down_sig,
                                                    'color_up_sig': UnivarA_Store.color_up_sig,}
        # Update to actually used parameters if possible
        if 'Test' in UnivarA_Store.current_univ_params:
            params_to_be_saved['Univariate Analysis']['univariate_test'] = UnivarA_Store.current_univ_params['Test']
            params_to_be_saved['Univariate Analysis']['expected_equal_variance'] = UnivarA_Store.current_univ_params[
                'Expected Equal Var.']
            params_to_be_saved['Univariate Analysis']['p_value_threshold'] = UnivarA_Store.current_univ_params['p-value']
            params_to_be_saved['Univariate Analysis']['fold_change_threshold'] = UnivarA_Store.current_univ_params[
                'Fold Change Threshold']


        # Saving Data Diversity Visualization Related Parameters
        params_to_be_saved['Data Visualization'] = {# Van Krevelen Related
                                                    'vk_highlight_by': dataviz_store.vk_highlight_by,
                                                    'vk_colour': dataviz_store.vk_colour,
                                                    'vk_size': dataviz_store.vk_size,
                                                    'vk_midpoint': dataviz_store.vk_midpoint,
                                                    'vk_max_dot_size': dataviz_store.vk_max_dot_size,
                                                    'vk_show_colorbar': dataviz_store.vk_show_colorbar,
                                                    'vk_draw_class_rectangle': dataviz_store.vk_draw_class_rectangle,
                                                    # Kendrick Mass Defect Related
                                                    'kmd_mass_rounding': dataviz_store.kmd_mass_rounding,
                                                    'kmd_max_dot_size': dataviz_store.kmd_max_dot_size,
                                                    # Chemical Composition Series Related
                                                    'ccs_bar_plot_type': dataviz_store.ccs_bar_plot_type,}


        # Saving BinSim Analysis Related Parameters
        params_to_be_saved['BinSim Analysis'] = {}

        # PCA
        params_to_be_saved['BinSim Analysis']['PCA'] = {'n_components': PCA_params_binsim.n_components, # Remember to load the widget too
                                                        'ellipse_draw': PCA_params_binsim.ellipse_draw,
                                                        'confidence': PCA_params_binsim.confidence,
                                                        'confidence_std': PCA_params_binsim.confidence_std,
                                                        'dot_size': PCA_params_binsim.dot_size}

        # HCA
        params_to_be_saved['BinSim Analysis']['HCA'] = {'dist_metric': HCA_params_binsim.dist_metric,
                                                        'link_metric': HCA_params_binsim.link_metric,
                                                        'fig_x': HCA_params_binsim.fig_x,
                                                        'fig_y': HCA_params_binsim.fig_y,
                                                        'dpi': HCA_params_binsim.dpi,}

        # PLS-DA
        params_to_be_saved['BinSim Analysis']['PLS-DA'] = {# Optimization Specific Related
                                                        'n_min_max_components': PLSDA_store_binsim.n_min_max_components,
                                                        # PLS-DA Model Fit
                                                        'n_fold': PLSDA_store_binsim.n_fold,
                                                        'scale': PLSDA_store_binsim.scale,
                                                        'n_components': PLSDA_store_binsim.n_components,
                                                        'n_iterations': PLSDA_store_binsim.n_iterations,
                                                        'metrics_to_use': PLSDA_store_binsim.metrics_to_use, # Technically can be different
                                                        'imp_feature_metric': PLSDA_store_binsim.imp_feature_metric,
                                                        # PLS Projection Related
                                                        'ellipse_draw': PLSDA_store_binsim.ellipse_draw,
                                                        'confidence': PLSDA_store_binsim.confidence,
                                                        'confidence_std': PLSDA_store_binsim.confidence_std,
                                                        'dot_size': PLSDA_store_binsim.dot_size,
                                                        # PLS-DA Permutation Test Related
                                                        'n_perm': PLSDA_store_binsim.n_perm,
                                                        'perm_metric': PLSDA_store_binsim.perm_metric,
                                                        'dpi': PLSDA_store_binsim.dpi, # Technically can be different
                                                        # PLS-DA ROC Curve Related
                                                        'roc_n_iter': PLSDA_store_binsim.roc_n_iter}
        # Update to actually used parameters if possible
        if 'n_folds' in PLSDA_store_binsim.current_plsda_params.keys():
            for k,v in PLSDA_store_binsim.current_plsda_params.items():
                if k == 'n_folds':
                    params_to_be_saved['BinSim Analysis']['PLS-DA']['n_fold'] = v
                elif k == 'feat_imp':
                    params_to_be_saved['BinSim Analysis']['PLS-DA']['imp_feature_metric'] = v
                else:
                    params_to_be_saved['BinSim Analysis']['PLS-DA'][k] = v
        if 'n_min_max_components' in PLSDA_store_binsim.current_other_plsda_params.keys():
            params_to_be_saved['BinSim Analysis']['PLS-DA']['n_min_max_components'] = PLSDA_store_binsim.current_other_plsda_params[
                'n_min_max_components']
        if 'ROC_filename' in PLSDA_store_binsim.current_other_plsda_params.keys():
            params_to_be_saved['BinSim Analysis']['PLS-DA']['roc_n_iter'] = PLSDA_store_binsim.current_other_plsda_params[
                'ROC_filename'].split('_')[5][:-10]
        if 'perm_metric' in PLSDA_store_binsim.current_plsda_params_permutation.keys():
            params_to_be_saved['BinSim Analysis']['PLS-DA'][
                'n_perm'] = PLSDA_store_binsim.current_plsda_params_permutation['n_permutations']
            params_to_be_saved['BinSim Analysis']['PLS-DA'][
                'perm_metric'] = PLSDA_store_binsim.current_plsda_params_permutation['perm_metric']

        # Random Forest
        params_to_be_saved['BinSim Analysis']['Random Forest'] = {# Optimization Specific Related
                                                            'n_min_max_trees': RF_store_binsim.n_min_max_trees,
                                                            'n_interval': RF_store_binsim.n_interval,
                                                            # Random Forest Model Fit
                                                            'n_fold': RF_store_binsim.n_fold,
                                                            'n_trees': RF_store_binsim.n_trees,
                                                            'n_iterations': RF_store_binsim.n_iterations,
                                                            'metrics_to_use': RF_store_binsim.metrics_to_use, # Technically can be different
                                                            # RF Permutation Test Related
                                                            'n_perm': RF_store_binsim.n_perm,
                                                            'perm_metric': RF_store_binsim.perm_metric,
                                                            'dpi': RF_store_binsim.dpi, # Technically can be different
                                                            # RF ROC Curve related
                                                            'roc_n_iter': RF_store_binsim.roc_n_iter,}
        # Update to actually used parameters if possible
        if 'n_folds' in RF_store_binsim.current_rf_params.keys():
            for k,v in RF_store_binsim.current_rf_params.items():
                if k == 'n_folds':
                    params_to_be_saved['BinSim Analysis']['Random Forest']['n_fold'] = v
                else:
                    params_to_be_saved['BinSim Analysis']['Random Forest'][k] = v
        if 'n_min_max_trees' in RF_store_binsim.current_other_rf_params.keys():
            params_to_be_saved['BinSim Analysis']['Random Forest']['n_min_max_trees'] = RF_store_binsim.current_other_rf_params[
                'n_min_max_trees']
            params_to_be_saved['BinSim Analysis']['Random Forest']['n_interval'] = RF_store_binsim.current_other_rf_params[
                'n_interval']
        if 'ROC_filename' in RF_store_binsim.current_other_rf_params.keys():
            params_to_be_saved['BinSim Analysis']['Random Forest']['roc_n_iter'] = RF_store_binsim.current_other_rf_params[
                'ROC_filename'].split('_')[5][:-10]
        if 'perm_metric' in RF_store_binsim.current_rf_params_permutation.keys():
            params_to_be_saved['BinSim Analysis']['Random Forest']['n_perm'] = RF_store_binsim.current_rf_params_permutation[
                'n_permutations']
            params_to_be_saved['BinSim Analysis']['Random Forest']['perm_metric'] = RF_store_binsim.current_rf_params_permutation[
                'perm_metric']


    # Save the parameters in a json file
    with open(filename+'.json', 'w') as f:
        json.dump(params_to_be_saved, f, indent=4)



# Loading parameters from saved json files.
def loading_parameters_in(params_to_load, filt_method, n_databases_show, n_databases, annotation_margin_method_radio,
                         annotation_ppm_deviation, annotation_Da_deviation, RepGen, adducts_to_search_widget, DB_dict,
                         PreTreatment_Method, checkbox_com_exc, com_exc_compounds, PCA_params, n_components_compute,
                         HCA_params, PLSDA_store, RF_store, UnivarA_Store, dataviz_store, PCA_params_binsim,
                         n_components_compute_binsim, HCA_params_binsim, PLSDA_store_binsim, RF_store_binsim,
                         params_pre_treat_loaded_in, params_analysis_loaded_in):
    "Load in Previously Saved Parameters for Data Analysis."

    # Data Pre-Processing and Pre-Treatment Related Parameters
    if params_pre_treat_loaded_in:
        # Loading Data Filtering Related Parameters
        filt_method.value = params_to_load['Data Filtering']['filt_method']
        #filt_kw

        # Loading Data Annotation Related Parameters
        # Main Parameters
        n_databases_show.value = params_to_load['Data Annotation']['n_databases']
        n_databases.value = params_to_load['Data Annotation']['n_databases']
        annotation_margin_method_radio.value = params_to_load['Data Annotation']['annotation_margin_method']
        annotation_ppm_deviation.value = params_to_load['Data Annotation']['annotation_ppm_deviation']
        annotation_Da_deviation.value = params_to_load['Data Annotation']['annotation_Da_deviation']
        if params_to_load['Data Reading']['type_of_mass_values_in_file'] == RepGen.type_of_mass_values_in_file:
            adducts_to_search_widget.value = params_to_load['Data Annotation']['adducts_to_be_searched']
        # Database Specific Parameters
        for d in params_to_load['Data Annotation'].keys():
            if d not in ['n_databases', 'annotation_margin_method', 'annotation_ppm_deviation', 'annotation_Da_deviation',
                         'adducts_to_be_searched']:
                DB_dict[str(int(d)+1)].content[0][1].value = params_to_load['Data Annotation'][d]["file"]
                DB_dict[str(int(d)+1)].content[1][1].value = params_to_load['Data Annotation'][d]["abv"]
                DB_dict[str(int(d)+1)].content[2][1].value = params_to_load['Data Annotation'][d]["IDcol"]
                DB_dict[str(int(d)+1)].content[3][1].value = params_to_load['Data Annotation'][d]["annotation"]
                DB_dict[str(int(d)+1)].content[4][1].value = params_to_load['Data Annotation'][d]["formula"]

        # Loading Data Pre-Treatment Related Parameters
        PreTreatment_Method.mvi_method = params_to_load['Data Pre-Treatment']['mvi_method']
        if params_to_load['Data Pre-Treatment']['mvi_method'] != 'Zero':
            PreTreatment_Method.mvi_kw = params_to_load['Data Pre-Treatment']['mvi_kw']
        #PreTreatment_Method.norm_method = params_to_load['Data Pre-Treatment']['norm_method']
        #PreTreatment_Method.norm_kw
        PreTreatment_Method.tf_method = params_to_load['Data Pre-Treatment']['tf_method']
        PreTreatment_Method.tf_kw = params_to_load['Data Pre-Treatment']['tf_kw']
        PreTreatment_Method.scaling_method = params_to_load['Data Pre-Treatment']['scaling_method']
        PreTreatment_Method.scaling_kw = params_to_load['Data Pre-Treatment']['scaling_kw']



    # Data Analysis Related Parameters
    if params_analysis_loaded_in:
        # Loading Common and Exclusive Compound Analysis Related Paramaeters
        if 'Com. and Exc. Comp.' in params_to_load.keys():
            # Stopping Figures from trying to be made
            com_exc_compounds.compute_fig = False
            # General
            checkbox_com_exc.value = params_to_load['Com. and Exc. Comp.']['chosen']
            # Venn Diagram
            com_exc_compounds.venn_alpha = params_to_load['Com. and Exc. Comp.']['venn_alpha']
            com_exc_compounds.controls.widgets['venn_alpha'].value = params_to_load['Com. and Exc. Comp.']['venn_alpha']
            com_exc_compounds.type_of_venn = params_to_load['Com. and Exc. Comp.']['type_of_venn']
            com_exc_compounds.controls.widgets['type_of_venn'].value = params_to_load['Com. and Exc. Comp.']['type_of_venn']
            com_exc_compounds.dpi_venn = params_to_load['Com. and Exc. Comp.']['dpi_venn']
            com_exc_compounds.controls.widgets['dpi_venn'].value = params_to_load['Com. and Exc. Comp.']['dpi_venn']
            # Intersection Plot Related
            if '%' in params_to_load['Com. and Exc. Comp.']['inter_include_counts_percentages']:
                com_exc_compounds.inter_include_counts_percentages = 'Show Nº and % of metabolites'
            elif 'metabolites' in params_to_load['Com. and Exc. Comp.']['inter_include_counts_percentages']:
                com_exc_compounds.inter_include_counts_percentages = 'Show Nº of metabolites'
            else:
                com_exc_compounds.inter_include_counts_percentages = 'Show neither'
            com_exc_compounds.dpi_inter = params_to_load['Com. and Exc. Comp.']['dpi_inter']
            com_exc_compounds.compute_fig = True

        # Loading Unsupervised Analysis Related Paramaeters
        # PCA
        if 'PCA' in params_to_load.keys():
            # Stopping Figures from trying to be made
            PCA_params.compute_fig = False

            PCA_params.n_components = params_to_load['PCA']['n_components']
            n_components_compute.value = params_to_load['PCA']['n_components']
            PCA_params.ellipse_draw = params_to_load['PCA']['ellipse_draw']
            PCA_params.confidence = params_to_load['PCA']['confidence']
            PCA_params.confidence_std = params_to_load['PCA']['confidence_std']
            PCA_params.dot_size = params_to_load['PCA']['dot_size']
            PCA_params.compute_fig = True

        # HCA
        if 'HCA' in params_to_load.keys():
            # Stopping Figures from trying to be made
            HCA_params.compute_fig = False
            HCA_params.dist_metric = params_to_load['HCA']['dist_metric']
            HCA_params.link_metric = params_to_load['HCA']['link_metric']
            HCA_params.fig_x = params_to_load['HCA']['fig_x']
            HCA_params.fig_y = params_to_load['HCA']['fig_y']
            HCA_params.dpi = params_to_load['HCA']['dpi']
            HCA_params.compute_fig = True

        # Loading Supervised Analysis Related Paramaeters
        # PLS-DA
        if 'PLS-DA' in params_to_load.keys():
            # Stopping Figures from trying to be made
            PLSDA_store.compute_fig = False
            # Optimization Specific Related
            PLSDA_store.n_min_max_components = tuple(params_to_load['PLS-DA']['n_min_max_components'])
            # PLS-DA Model Fit
            PLSDA_store.n_fold = params_to_load['PLS-DA']['n_fold']
            PLSDA_store.scale = params_to_load['PLS-DA']['scale']
            PLSDA_store.n_components = params_to_load['PLS-DA']['n_components']
            PLSDA_store.n_iterations = params_to_load['PLS-DA']['n_iterations']
            PLSDA_store.metrics_to_use = params_to_load['PLS-DA']['metrics_to_use']
            PLSDA_store.imp_feature_metric = params_to_load['PLS-DA']['imp_feature_metric']
            # PLS Projection Related
            PLSDA_store.param['ellipse_draw'].default = params_to_load['PLS-DA']['ellipse_draw']
            PLSDA_store.ellipse_draw = params_to_load['PLS-DA']['ellipse_draw']
            PLSDA_store.controls_projection.widgets['ellipse_draw'].value = params_to_load['PLS-DA']['ellipse_draw']
            PLSDA_store.param['confidence'].default = params_to_load['PLS-DA']['confidence']
            PLSDA_store.confidence = params_to_load['PLS-DA']['confidence']
            PLSDA_store.controls_projection.widgets['confidence'].value = params_to_load['PLS-DA']['confidence']
            PLSDA_store.param['confidence_std'].default = params_to_load['PLS-DA']['confidence_std']
            PLSDA_store.confidence_std = params_to_load['PLS-DA']['confidence_std']
            PLSDA_store.controls_projection.widgets['confidence_std'].value = params_to_load['PLS-DA']['confidence_std']
            PLSDA_store.param['dot_size'].default = params_to_load['PLS-DA']['dot_size']
            PLSDA_store.dot_size = params_to_load['PLS-DA']['dot_size']
            PLSDA_store.controls_projection.widgets['dot_size'].value = params_to_load['PLS-DA']['dot_size']
            # PLS-DA Permutation Test Related
            PLSDA_store.n_perm = params_to_load['PLS-DA']['n_perm']
            PLSDA_store.perm_metric = params_to_load['PLS-DA']['perm_metric']
            PLSDA_store.dpi = params_to_load['PLS-DA']['dpi']
            # PLS-DA ROC Curve Related
            PLSDA_store.roc_n_iter = int(params_to_load['PLS-DA']['roc_n_iter'])
            PLSDA_store.compute_fig = True

        # Random Forest
        if 'Random Forest' in params_to_load.keys():
            # Optimization Specific Related
            RF_store.n_min_max_trees = tuple(params_to_load['Random Forest']['n_min_max_trees'])
            RF_store.n_interval = params_to_load['Random Forest']['n_interval']
            # Random Forest Model Fit
            RF_store.n_fold = params_to_load['Random Forest']['n_fold']
            RF_store.n_trees = params_to_load['Random Forest']['n_trees']
            RF_store.n_iterations = params_to_load['Random Forest']['n_iterations']
            RF_store.metrics_to_use = params_to_load['Random Forest']['metrics_to_use']
            # RF Permutation Test Related
            RF_store.n_perm = params_to_load['Random Forest']['n_perm']
            RF_store.perm_metric = params_to_load['Random Forest']['perm_metric']
            RF_store.dpi = params_to_load['Random Forest']['dpi']
            # RF ROC Curve Related
            RF_store.roc_n_iter = int(params_to_load['Random Forest']['roc_n_iter'])

        # Loading Univariate Analysis Related Parameters
        if 'Univariate Analysis' in params_to_load.keys():
            # Stopping Figures from trying to be made
            UnivarA_Store.compute_fig = False
            UnivarA_Store.univariate_test = params_to_load['Univariate Analysis']['univariate_test']
            UnivarA_Store.expected_equal_variance = params_to_load['Univariate Analysis']['expected_equal_variance']
            UnivarA_Store.p_value_threshold = params_to_load['Univariate Analysis']['p_value_threshold']
            UnivarA_Store.fold_change_threshold = params_to_load['Univariate Analysis']['fold_change_threshold']
            UnivarA_Store.color_non_sig = params_to_load['Univariate Analysis']['color_non_sig']
            UnivarA_Store.color_down_sig = params_to_load['Univariate Analysis']['color_down_sig']
            UnivarA_Store.color_up_sig = params_to_load['Univariate Analysis']['color_up_sig']
            UnivarA_Store.compute_fig = True

        # Loading Data Diversity Visualization Related Parameters:
        if 'Data Visualization' in params_to_load.keys():
            # Stopping Figures from trying to be made
            dataviz_store.compute_fig = False
            # Van Krevelen Related
            dataviz_store.vk_highlight_by = params_to_load['Data Visualization']['vk_highlight_by']
            dataviz_store.vk_colour = params_to_load['Data Visualization']['vk_colour']
            dataviz_store.vk_size = params_to_load['Data Visualization']['vk_size']
            dataviz_store.vk_midpoint = params_to_load['Data Visualization']['vk_midpoint']
            dataviz_store.vk_max_dot_size = params_to_load['Data Visualization']['vk_max_dot_size']
            dataviz_store.vk_show_colorbar = params_to_load['Data Visualization']['vk_show_colorbar']
            dataviz_store.vk_draw_class_rectangle = params_to_load['Data Visualization']['vk_draw_class_rectangle']
            # Kendrick Mass Defect Related
            dataviz_store.kmd_mass_rounding = params_to_load['Data Visualization']['kmd_mass_rounding']
            dataviz_store.kmd_max_dot_size = params_to_load['Data Visualization']['kmd_max_dot_size']
            # Chemical Composition Series Related
            dataviz_store.ccs_bar_plot_type = params_to_load['Data Visualization']['ccs_bar_plot_type']
            dataviz_store.compute_fig = True

        # Loading BinSim Analysis Related Parameters
        if 'BinSim Analysis' in params_to_load.keys():

            # PCA
            # Stopping Figures from trying to be made
            PCA_params_binsim.compute_fig = False
            PCA_params_binsim.n_components = params_to_load['BinSim Analysis']['PCA']['n_components']
            n_components_compute_binsim.value = params_to_load['BinSim Analysis']['PCA']['n_components']
            PCA_params_binsim.ellipse_draw = params_to_load['BinSim Analysis']['PCA']['ellipse_draw']
            PCA_params_binsim.confidence = params_to_load['BinSim Analysis']['PCA']['confidence']
            PCA_params_binsim.confidence_std = params_to_load['BinSim Analysis']['PCA']['confidence_std']
            PCA_params_binsim.dot_size = params_to_load['BinSim Analysis']['PCA']['dot_size']
            PCA_params_binsim.compute_fig = True

            # HCA
            # Stopping Figures from trying to be made
            HCA_params_binsim.compute_fig = False
            #HCA_params_binsim.dist_metric = params_to_load['BinSim Analysis']['HCA']['dist_metric']
            HCA_params_binsim.link_metric = params_to_load['BinSim Analysis']['HCA']['link_metric']
            HCA_params_binsim.fig_x = params_to_load['BinSim Analysis']['HCA']['fig_x']
            HCA_params_binsim.fig_y = params_to_load['BinSim Analysis']['HCA']['fig_y']
            HCA_params_binsim.dpi = params_to_load['BinSim Analysis']['HCA']['dpi']
            HCA_params_binsim.compute_fig = True

            # PLS-DA
            # Stopping Figures from trying to be made
            PLSDA_store_binsim.compute_fig = False
            # Optimization Specific Related
            PLSDA_store_binsim.n_min_max_components = tuple(params_to_load['BinSim Analysis']['PLS-DA']['n_min_max_components'])
            # PLS-DA Model Fit
            PLSDA_store_binsim.n_fold = params_to_load['BinSim Analysis']['PLS-DA']['n_fold']
            PLSDA_store_binsim.scale = params_to_load['BinSim Analysis']['PLS-DA']['scale']
            PLSDA_store_binsim.n_components = params_to_load['BinSim Analysis']['PLS-DA']['n_components']
            PLSDA_store_binsim.n_iterations = params_to_load['BinSim Analysis']['PLS-DA']['n_iterations']
            PLSDA_store_binsim.metrics_to_use = params_to_load['BinSim Analysis']['PLS-DA']['metrics_to_use']
            PLSDA_store_binsim.imp_feature_metric = params_to_load['BinSim Analysis']['PLS-DA']['imp_feature_metric']
            # PLS Projection Related
            PLSDA_store_binsim.param['ellipse_draw'].default = params_to_load['BinSim Analysis']['PLS-DA']['ellipse_draw']
            PLSDA_store_binsim.ellipse_draw = params_to_load['BinSim Analysis']['PLS-DA']['ellipse_draw']
            PLSDA_store_binsim.controls_projection.widgets['ellipse_draw'].value = params_to_load[
                'BinSim Analysis']['PLS-DA']['ellipse_draw']
            PLSDA_store_binsim.param['confidence'].default = params_to_load['BinSim Analysis']['PLS-DA']['confidence']
            PLSDA_store_binsim.confidence = params_to_load['BinSim Analysis']['PLS-DA']['confidence']
            PLSDA_store_binsim.controls_projection.widgets['confidence'].value = params_to_load[
                'BinSim Analysis']['PLS-DA']['confidence']
            PLSDA_store_binsim.param['confidence_std'].default = params_to_load['BinSim Analysis']['PLS-DA']['confidence_std']
            PLSDA_store_binsim.confidence_std = params_to_load['BinSim Analysis']['PLS-DA']['confidence_std']
            PLSDA_store_binsim.controls_projection.widgets['confidence_std'].value = params_to_load[
                'BinSim Analysis']['PLS-DA']['confidence_std']
            PLSDA_store_binsim.param['dot_size'].default = params_to_load['BinSim Analysis']['PLS-DA']['dot_size']
            PLSDA_store_binsim.dot_size = params_to_load['BinSim Analysis']['PLS-DA']['dot_size']
            PLSDA_store_binsim.controls_projection.widgets['dot_size'].value = params_to_load[
                'BinSim Analysis']['PLS-DA']['dot_size']
            # PLS-DA Permutation Test Related
            PLSDA_store_binsim.n_perm = params_to_load['BinSim Analysis']['PLS-DA']['n_perm']
            PLSDA_store_binsim.perm_metric = params_to_load['BinSim Analysis']['PLS-DA']['perm_metric']
            PLSDA_store_binsim.dpi = params_to_load['BinSim Analysis']['PLS-DA']['dpi']
            # PLS-DA ROC Curve Related
            PLSDA_store_binsim.roc_n_iter = int(params_to_load['BinSim Analysis']['PLS-DA']['roc_n_iter'])
            PLSDA_store_binsim.compute_fig = True

            # Random Forest
            # Optimization Specific Related
            RF_store_binsim.n_min_max_trees = tuple(params_to_load['BinSim Analysis']['Random Forest']['n_min_max_trees'])
            RF_store_binsim.n_interval = params_to_load['BinSim Analysis']['Random Forest']['n_interval']
            # Random Forest Model Fit
            RF_store_binsim.n_fold = params_to_load['BinSim Analysis']['Random Forest']['n_fold']
            RF_store_binsim.n_trees = params_to_load['BinSim Analysis']['Random Forest']['n_trees']
            RF_store_binsim.n_iterations = params_to_load['BinSim Analysis']['Random Forest']['n_iterations']
            RF_store_binsim.metrics_to_use = params_to_load['BinSim Analysis']['Random Forest']['metrics_to_use']
            # RF Permutation Test Related
            RF_store_binsim.n_perm = params_to_load['BinSim Analysis']['Random Forest']['n_perm']
            RF_store_binsim.perm_metric = params_to_load['BinSim Analysis']['Random Forest']['perm_metric']
            RF_store_binsim.dpi = params_to_load['BinSim Analysis']['Random Forest']['dpi']
            # RF ROC Curve Related
            RF_store_binsim.roc_n_iter = int(params_to_load['BinSim Analysis']['Random Forest']['roc_n_iter'])
